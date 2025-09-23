"""
Обработчики колбеков для CyberKitty Transkribator
"""

import json
from datetime import datetime
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ContextTypes

from transkribator_modules.config import logger
from transkribator_modules.db.database import SessionLocal, UserService, ApiKeyService, PromoCodeService
from transkribator_modules.db.models import ApiKey, PlanType
from transkribator_modules.bot.payments import handle_payment_callback, show_payment_plans, initiate_payment, initiate_yukassa_payment

async def handle_callback_query(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обрабатывает колбек запросы от кнопок."""
    query = update.callback_query
    await query.answer()

    data = query.data
    logger.info(f"Получен колбек: {data}")
    logger.info(f"Полный update: {update.to_dict() if hasattr(update, 'to_dict') else str(update)}")

    # Обработка различных типов колбеков
    if data == "show_payment_plans":
        logger.info("Получен колбек show_payment_plans")
        await show_payment_plans(update, context)

    elif data == "personal_cabinet":
        await show_personal_cabinet(update, context)

    elif data == "show_help":
        from transkribator_modules.bot.commands import help_command
        await help_command(update, context)

    elif data.startswith("buy_plan_"):
        if data.endswith("_stars"):
            # Платеж через Telegram Stars
            plan_id = data.replace("buy_plan_", "").replace("_stars", "")
            logger.info(f"Обнаружен колбек оплаты плана через Stars: {data}, извлеченный plan_id: {plan_id}")
            await initiate_payment(update, context, plan_id)
        elif data.endswith("_yukassa"):
            # Платеж через ЮКассу
            plan_id = data.replace("buy_plan_", "").replace("_yukassa", "")
            logger.info(f"Обнаружен колбек оплаты плана через ЮКассу: {data}, извлеченный plan_id: {plan_id}")
            await initiate_yukassa_payment(update, context, plan_id)
        else:
            # Старый формат для обратной совместимости
            plan_id = data.replace("buy_plan_", "")
            logger.info(f"Обнаружен колбек оплаты плана (старый формат): {data}, извлеченный plan_id: {plan_id}")
            await initiate_payment(update, context, plan_id)


    elif data == "show_stats":
        from transkribator_modules.bot.commands import stats_command
        await stats_command(update, context)

    elif data == "show_api_keys":
        await show_api_keys(update, context)

    elif data == "enter_promo_code":
        await enter_promo_code(update, context)

    elif data == "show_promo_codes":
        from transkribator_modules.bot.commands import promo_codes_command
        await promo_codes_command(update, context)

    elif data == "show_plans":
        await show_plans_callback(query, update.effective_user)

    elif data == "create_api_key":
        await create_api_key_callback(query, update.effective_user)

    elif data == "list_api_keys":
        await list_api_keys_callback(query, update.effective_user)

    elif data.startswith("delete_api_key_"):
        key_id = int(data.split("_")[-1])
        await delete_api_key_callback(query, update.effective_user, key_id)

    elif data == "back_to_start":
        await back_to_start_callback(query, update.effective_user)

    elif data.startswith("brief_summary_") or data.startswith("detailed_summary_"):
        await handle_summary_callback(update, context)

    elif data.startswith("process_transcript_"):
        await handle_process_transcript_callback(update, context)

    elif data.startswith("send_more_"):
        await handle_send_more_callback(update, context)

    elif data.startswith("main_menu_"):
        await handle_main_menu_callback(update, context)

    else:
        await query.edit_message_text("Неизвестная команда")

async def show_personal_cabinet(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Показывает личный кабинет пользователя."""
    try:
        user = update.effective_user
        db = SessionLocal()

        try:
            user_service = UserService(db)
            from transkribator_modules.db.database import TranscriptionService
            transcription_service = TranscriptionService(db)

            db_user = user_service.get_or_create_user(telegram_id=user.id)
            usage_info = user_service.get_usage_info(db_user)

            # Получаем количество обработанных файлов
            transcriptions_count = transcription_service.get_user_transcriptions_count(db_user)

            # Определяем статус тарифа
            plan_status = ""
            if db_user.plan_expires_at:
                if db_user.plan_expires_at > datetime.utcnow():
                    days_left = (db_user.plan_expires_at - datetime.utcnow()).days
                    plan_status = f"(истекает через {days_left} дн.)"
                else:
                    plan_status = "(истек)"
            elif db_user.current_plan != "free":
                plan_status = "(бессрочно 🎉)"

            # Формируем информацию об использовании
            usage_text = ""
            if usage_info['current_plan'] == 'free':
                remaining = usage_info['generations_remaining']
                percentage = usage_info['usage_percentage']
                usage_text = f"""📊 **Использование в этом месяце:**
• Использовано: {usage_info['generations_used_this_month']} из {usage_info['generations_limit']} генераций
• Осталось: {remaining} генераций
• Всего генераций: {usage_info['total_generations']}
• Минут транскрибировано: {usage_info['minutes_used_this_month']:.1f} мин"""
            elif usage_info['minutes_limit']:
                remaining = usage_info['minutes_remaining']
                percentage = usage_info['usage_percentage']
                usage_text = f"""📊 **Использование в этом месяце:**
• Использовано: {usage_info['minutes_used_this_month']:.1f} из {usage_info['minutes_limit']:.0f} мин
• Осталось: {remaining:.1f} мин"""
            else:
                usage_text = f"""📊 **Использование в этом месяце:**
• Использовано: {usage_info['minutes_used_this_month']:.1f} мин
• Лимит: Безлимитно ♾️"""

            cabinet_text = f"""🐱 **Личный кабинет**

👤 **Профиль:**
• Имя: {user.first_name or 'Котик'} {user.last_name or ''}
• План: {usage_info['plan_display_name']} {plan_status}

{usage_text}

📈 **Статистика:**
• Файлов обработано: {transcriptions_count}
• Всего транскрибировано: {usage_info['total_minutes_transcribed']:.1f} мин
• Последняя активность: {db_user.updated_at.strftime('%d.%m.%Y %H:%M') if db_user.updated_at else 'Никогда'}

**Доступные функции:**
• Транскрипция видео и аудио
• Обработка файлов до 2 ГБ
• Техническая поддержка

Для расширения возможностей рассмотрите PRO подписку! 🚀"""
        finally:
            db.close()

        keyboard = [
            [InlineKeyboardButton("💎 Тарифы", callback_data="show_payment_plans")],
            [InlineKeyboardButton("🎁 Промокоды", callback_data="enter_promo_code")],
            [InlineKeyboardButton("❓ Помощь", callback_data="show_help")]
        ]

        reply_markup = InlineKeyboardMarkup(keyboard)

        if update.callback_query:
            await update.callback_query.edit_message_text(
                cabinet_text, reply_markup=reply_markup, parse_mode='Markdown'
            )
        else:
            await update.message.reply_text(
                cabinet_text, reply_markup=reply_markup, parse_mode='Markdown'
            )

    except Exception as e:
        logger.error(f"Ошибка в личном кабинете: {e}")
        await update.message.reply_text("❌ Ошибка при загрузке личного кабинета")

async def show_plans_callback(query, user):
    """Показать тарифные планы"""
    from transkribator_modules.db.database import get_plans

    plans = get_plans()
    plans_text = "💰 **Тарифные планы:**\n\n"

    for plan in plans:
        features = []
        if plan.features:
            try:
                features = json.loads(plan.features)
            except:
                features = [plan.features]

        minutes_text = f"{plan.minutes_per_month:.0f} минут" if plan.minutes_per_month else "Безлимитно"
        price_text = f"{plan.price_rub:.0f} ₽" if plan.price_rub > 0 else "Бесплатно"

        plans_text += f"**{plan.display_name}** - {price_text}\n"
        plans_text += f"• {minutes_text} в месяц\n"
        plans_text += f"• Файлы до {plan.max_file_size_mb:.0f} МБ\n"

        for feature in features:
            plans_text += f"• {feature}\n"

        plans_text += f"_{plan.description}_\n\n"

    plans_text += "⭐ **Покупка через Telegram Stars**"

    keyboard = [
        [InlineKeyboardButton("⭐ Купить план", callback_data="show_payment_plans")],
        [InlineKeyboardButton("🔙 Личный кабинет", callback_data="personal_cabinet")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await query.edit_message_text(plans_text, reply_markup=reply_markup, parse_mode='Markdown')

async def show_stats_callback(query, user):
    """Показать статистику пользователя"""
    db = SessionLocal()
    try:
        user_service = UserService(db)
        from transkribator_modules.db.database import TranscriptionService
        transcription_service = TranscriptionService(db)

        db_user = user_service.get_or_create_user(telegram_id=user.id)
        usage_info = user_service.get_usage_info(db_user)

        recent_transcriptions = transcription_service.get_user_transcriptions(db_user, limit=5)

        stats_text = f"""📊 **Ваша статистика:**

👤 **Профиль:**
• Telegram ID: `{user.id}`
• План: {usage_info['plan_display_name']}
• Зарегистрирован: {db_user.created_at.strftime('%d.%m.%Y')}

📈 **Использование:**"""

        # Для бесплатного тарифа показываем генерации
        if usage_info['current_plan'] == 'free':
            remaining = usage_info['generations_remaining']
            percentage = usage_info['usage_percentage']
            stats_text += f"\n• В этом месяце: {usage_info['generations_used_this_month']} генераций"
            stats_text += f"\n• Лимит: {usage_info['generations_limit']} генераций"
            stats_text += f"\n• Осталось: {remaining} генераций ({100-percentage:.1f}%)"
            stats_text += f"\n• Всего генераций: {usage_info['total_generations']}"
        elif usage_info['minutes_limit']:
            remaining = usage_info['minutes_remaining']
            percentage = usage_info['usage_percentage']
            stats_text += f"\n• В этом месяце: {usage_info['minutes_used_this_month']:.1f} мин"
            stats_text += f"\n• Лимит: {usage_info['minutes_limit']:.0f} мин"
            stats_text += f"\n• Осталось: {remaining:.1f} мин ({100-percentage:.1f}%)"
        else:
            stats_text += f"\n• В этом месяце: {usage_info['minutes_used_this_month']:.1f} мин"
            stats_text += f"\n• Лимит: Безлимитно ♾️"

        stats_text += f"\n• Всего транскрибировано: {usage_info['total_minutes_transcribed']:.1f} мин"

        if recent_transcriptions:
            stats_text += f"\n\n🎬 **Последние транскрибации:**"
            for i, trans in enumerate(recent_transcriptions, 1):
                date_str = trans.created_at.strftime('%d.%m %H:%M')
                stats_text += f"\n{i}. {trans.filename or 'Видео'} ({trans.audio_duration_minutes:.1f} мин) - {date_str}"

        keyboard = [
            [InlineKeyboardButton("📊 Планы", callback_data="show_plans")],
            [InlineKeyboardButton("🔙 Назад", callback_data="back_to_start")]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)

        await query.edit_message_text(stats_text, reply_markup=reply_markup, parse_mode='Markdown')

    finally:
        db.close()

async def show_api_keys_callback(query, user):
    """Показать API ключи пользователя"""
    db = SessionLocal()
    try:
        user_service = UserService(db)
        db_user = user_service.get_or_create_user(telegram_id=user.id)
        plan = user_service.get_user_plan(db_user)

        # Проверяем доступ к API
        if not plan or plan.name == PlanType.FREE:
            api_text = f"""🔑 **API доступ**

❌ API доступ недоступен для плана "{plan.display_name}"

API доступен начиная с плана "Профессиональный".

💡 Обновите план для получения доступа к API."""

            keyboard = [
                [InlineKeyboardButton("📊 Посмотреть планы", callback_data="show_plans")],
                [InlineKeyboardButton("🔙 Назад", callback_data="back_to_start")]
            ]
        else:
            # Получаем API ключи
            api_keys = db.query(ApiKey).filter(
                ApiKey.user_id == db_user.id,
                ApiKey.is_active == True
            ).all()

            api_text = f"""🔑 **Управление API ключами**

✅ API доступ активен для плана "{plan.display_name}"

**Ваши API ключи:** ({len(api_keys)}/5)"""

            if api_keys:
                for i, key in enumerate(api_keys, 1):
                    last_used = key.last_used_at.strftime('%d.%m.%Y') if key.last_used_at else "Не использовался"
                    limit_text = f"{key.minutes_limit:.0f} мин" if key.minutes_limit else "Безлимитно"
                    api_text += f"\n{i}. {key.name}"
                    api_text += f"\n   • Лимит: {limit_text}"
                    api_text += f"\n   • Использовано: {key.minutes_used:.1f} мин"
                    api_text += f"\n   • Последнее использование: {last_used}"
            else:
                api_text += "\n\nУ вас пока нет API ключей."

            api_text += f"\n\n📖 **Документация API:**"
            api_text += f"\nБазовый URL: `http://localhost:8000`"
            api_text += f"\nЗаголовок: `Authorization: Bearer YOUR_API_KEY`"
            api_text += f"\nЭндпоинт: `POST /transcribe`"

            keyboard = []
            if len(api_keys) < 5:  # Лимит 5 ключей
                keyboard.append([InlineKeyboardButton("➕ Создать ключ", callback_data="create_api_key")])

            if api_keys:
                keyboard.append([InlineKeyboardButton("📋 Управление ключами", callback_data="list_api_keys")])

            keyboard.append([InlineKeyboardButton("🔙 Назад", callback_data="back_to_start")])

        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(api_text, reply_markup=reply_markup, parse_mode='Markdown')

    finally:
        db.close()

async def create_api_key_callback(query, user):
    """Создать новый API ключ"""
    db = SessionLocal()
    try:
        user_service = UserService(db)
        api_key_service = ApiKeyService(db)

        db_user = user_service.get_or_create_user(telegram_id=user.id)
        plan = user_service.get_user_plan(db_user)

        # Проверяем лимиты
        existing_keys = db.query(ApiKey).filter(
            ApiKey.user_id == db_user.id,
            ApiKey.is_active == True
        ).count()

        if existing_keys >= 5:
            await query.edit_message_text(
                "❌ Достигнут лимит API ключей (5 штук). Удалите неиспользуемые ключи.",
                reply_markup=InlineKeyboardMarkup([[
                    InlineKeyboardButton("🔙 Назад", callback_data="show_api_keys")
                ]])
            )
            return

        # Создаем ключ
        raw_key, api_key = api_key_service.generate_api_key(
            user=db_user,
            name=f"API Key {existing_keys + 1}"
        )

        success_text = f"""✅ **API ключ создан!**

🔑 **Ваш новый API ключ:**
```
{raw_key}
```

⚠️ **ВАЖНО:** Сохраните этот ключ в безопасном месте! Он больше не будет показан.

📖 **Пример использования:**
```bash
curl -X POST "http://localhost:8000/transcribe" \\
  -H "Authorization: Bearer {raw_key}" \\
  -F "file=@video.mp4"
```

🔒 **Безопасность:**
• Не передавайте ключ третьим лицам
• Используйте HTTPS в продакшене
• Регулярно обновляйте ключи"""

        keyboard = [
            [InlineKeyboardButton("📋 Мои ключи", callback_data="show_api_keys")],
            [InlineKeyboardButton("🔙 Главное меню", callback_data="back_to_start")]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)

        await query.edit_message_text(success_text, reply_markup=reply_markup, parse_mode='Markdown')

    except Exception as e:
        logger.error(f"Ошибка при создании API ключа: {e}")
        await query.edit_message_text(
            "Произошла ошибка при создании API ключа. *смущенно прячет мордочку*",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("🔙 Назад", callback_data="show_api_keys")
            ]])
        )
    finally:
        db.close()

async def list_api_keys_callback(query, user):
    """Показать список API ключей для управления"""
    db = SessionLocal()
    try:
        user_service = UserService(db)
        db_user = user_service.get_or_create_user(telegram_id=user.id)

        api_keys = db.query(ApiKey).filter(
            ApiKey.user_id == db_user.id,
            ApiKey.is_active == True
        ).all()

        if not api_keys:
            await query.edit_message_text(
                "У вас нет активных API ключей.",
                reply_markup=InlineKeyboardMarkup([[
                    InlineKeyboardButton("🔙 Назад", callback_data="show_api_keys")
                ]])
            )
            return

        keys_text = "🔑 **Управление API ключами:**\n\n"
        keyboard = []

        for i, key in enumerate(api_keys, 1):
            created = key.created_at.strftime('%d.%m.%Y')
            last_used = key.last_used_at.strftime('%d.%m.%Y') if key.last_used_at else "Не использовался"
            limit_text = f"{key.minutes_limit:.0f} мин" if key.minutes_limit else "Безлимитно"

            keys_text += f"**{i}. {key.name}**\n"
            keys_text += f"• Создан: {created}\n"
            keys_text += f"• Лимит: {limit_text}\n"
            keys_text += f"• Использовано: {key.minutes_used:.1f} мин\n"
            keys_text += f"• Последнее использование: {last_used}\n\n"

            # Добавляем кнопку удаления
            keyboard.append([InlineKeyboardButton(
                f"🗑 Удалить '{key.name}'",
                callback_data=f"delete_api_key_{key.id}"
            )])

        keyboard.append([InlineKeyboardButton("🔙 Назад", callback_data="show_api_keys")])
        reply_markup = InlineKeyboardMarkup(keyboard)

        await query.edit_message_text(keys_text, reply_markup=reply_markup, parse_mode='Markdown')

    finally:
        db.close()

async def delete_api_key_callback(query, user, key_id):
    """Удалить API ключ"""
    db = SessionLocal()
    try:
        user_service = UserService(db)
        db_user = user_service.get_or_create_user(telegram_id=user.id)

        # Находим ключ
        api_key = db.query(ApiKey).filter(
            ApiKey.id == key_id,
            ApiKey.user_id == db_user.id,
            ApiKey.is_active == True
        ).first()

        if not api_key:
            await query.edit_message_text(
                "API ключ не найден или уже удален.",
                reply_markup=InlineKeyboardMarkup([[
                    InlineKeyboardButton("🔙 Назад", callback_data="list_api_keys")
                ]])
            )
            return

        # Деактивируем ключ
        api_key.is_active = False
        db.commit()

        await query.edit_message_text(
            f"✅ API ключ '{api_key.name}' успешно удален.",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("📋 Мои ключи", callback_data="show_api_keys"),
                InlineKeyboardButton("🔙 Главное меню", callback_data="back_to_start")
            ]])
        )

    except Exception as e:
        logger.error(f"Ошибка при удалении API ключа: {e}")
        await query.edit_message_text(
            "Произошла ошибка при удалении ключа.",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("🔙 Назад", callback_data="list_api_keys")
            ]])
        )
    finally:
        db.close()

async def back_to_start_callback(query, user):
    """Вернуться в главное меню"""
    welcome_text = f"""🐱 **Мяу! Добро пожаловать в Cyberkitty19 Transkribator!**

Привет, {user.first_name or 'котик'}! Я умный котик-транскрибатор, который превращает твои видео в текст!

🎬 **Что я умею:**
• Транскрибирую видео любого формата в текст
• Форматирую текст с помощью ИИ
• Создаю краткие и подробные саммари
• Работаю с большими файлами через API

🚀 **Как это работает:**
Просто отправь мне видео, и я создам красивую текстовую расшифровку! Можешь выбрать обычную транскрибацию или с ИИ-форматированием.

💡 **Готов начать?**
Нажми кнопку ниже, чтобы войти в личный кабинет, или просто отправь мне видео!

*мурчит и виляет хвостиком* 🐾"""

    keyboard = [
        [InlineKeyboardButton("🏠 Личный кабинет", callback_data="personal_cabinet")],
        [InlineKeyboardButton("💡 Помощь", callback_data="show_help")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await query.edit_message_text(welcome_text, reply_markup=reply_markup, parse_mode='Markdown')

async def show_api_keys(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Показывает API ключи пользователя."""
    try:
        api_text = """🔑 **API ключи**

🚧 **API находится в разработке**

**Планируемые возможности:**
• REST API для транскрипции
• Webhook уведомления
• Интеграция с внешними сервисами
• Пакетная обработка файлов

**Требования:**
• PRO подписка или выше
• Подтвержденный аккаунт
• Согласие с условиями использования

Следите за обновлениями! 🚀"""

        from telegram import InlineKeyboardButton, InlineKeyboardMarkup

        keyboard = [
            [InlineKeyboardButton("🔙 Личный кабинет", callback_data="personal_cabinet")]
        ]

        reply_markup = InlineKeyboardMarkup(keyboard)

        await update.callback_query.edit_message_text(
            api_text, reply_markup=reply_markup, parse_mode='Markdown'
        )

    except Exception as e:
        logger.error(f"Ошибка при показе API ключей: {e}")
        await update.callback_query.edit_message_text("❌ Ошибка при загрузке API ключей")

async def enter_promo_code(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Предлагает ввести промокод."""
    try:
        promo_text = """🎁 **Ввод промокода**

Отправьте промокод в следующем сообщении или используйте команду:
`/promo [ваш_промокод]`

**Примеры:**
• `/promo WELCOME10`
• `/promo PREMIUM30`

**Где найти промокоды:**
• Официальный канал разработчика
• Социальные сети
• Специальные акции

Ждем ваш промокод! 🔥"""

        from telegram import InlineKeyboardButton, InlineKeyboardMarkup

        keyboard = [
            [InlineKeyboardButton("🔙 Личный кабинет", callback_data="personal_cabinet")]
        ]

        reply_markup = InlineKeyboardMarkup(keyboard)

        await update.callback_query.edit_message_text(
            promo_text, reply_markup=reply_markup, parse_mode='Markdown'
        )

    except Exception as e:
        logger.error(f"Ошибка при вводе промокода: {e}")
        await update.callback_query.edit_message_text("❌ Ошибка при обработке промокода")

async def handle_summary_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обрабатывает запросы на создание саммари."""
    query = update.callback_query
    await query.answer()

    try:
        data = query.data
        user_id = update.effective_user.id

        # Определяем тип саммари
        if data.startswith("brief_summary_"):
            summary_type = "brief"
            summary_type_ru = "краткое"
        elif data.startswith("detailed_summary_"):
            summary_type = "detailed"
            summary_type_ru = "подробное"
        else:
            await query.edit_message_text("❌ Неизвестный тип саммари")
            return

        # Обновляем сообщение
        await query.edit_message_text(
            f"🤖 Создаю {summary_type_ru} саммари...\n\n"
            f"*сосредоточенно работает*\n"
            f"Это может занять некоторое время...",
            parse_mode='Markdown'
        )

        # Получаем последнюю транскрипцию пользователя из базы данных
        db = SessionLocal()
        try:
            from transkribator_modules.db.database import TranscriptionService
            transcription_service = TranscriptionService(db)

            # Получаем пользователя
            user_service = UserService(db)
            user = user_service.get_or_create_user(telegram_id=user_id)

            # Получаем последнюю транскрипцию пользователя
            transcriptions = transcription_service.get_user_transcriptions(user, limit=1)

            if not transcriptions:
                await query.edit_message_text(
                    "❌ Не найдено транскрипций для создания саммари.\n\n"
                    "Сначала отправьте файл для транскрипции!"
                )
                return

            latest_transcription = transcriptions[0]
            transcript_text = latest_transcription.formatted_transcript or latest_transcription.raw_transcript

            if not transcript_text:
                await query.edit_message_text("❌ Транскрипция пуста")
                return

            # Импортируем функции генерации саммари
            from transkribator_modules.transcribe.transcriber_v4 import generate_detailed_summary, generate_brief_summary

            # Генерируем саммари
            if summary_type == "brief":
                summary = await generate_brief_summary(transcript_text)
            else:
                summary = await generate_detailed_summary(transcript_text)

            if not summary:
                await query.edit_message_text(
                    "❌ Не удалось создать саммари.\n\n"
                    "Возможно, сервис временно недоступен. Попробуйте позже."
                )
                return

            # Отправляем саммари
            summary_text = f"📋 **{summary_type_ru.title()} саммари:**\n\n{summary}\n\n@CyberKitty19_bot"

            # Если саммари длинное, отправляем файлом
            if len(summary_text) > 4000:
                from docx import Document
                from pathlib import Path

                docx_path = Path("/tmp") / f"summary_{user_id}_{summary_type}.docx"
                document = Document()
                document.add_heading(f"{summary_type_ru.title()} саммари", 0)
                document.add_paragraph(summary)
                document.save(docx_path)

                with open(docx_path, 'rb') as f:
                    await query.message.reply_document(
                        document=f,
                        filename=f"summary_{summary_type}.docx",
                        caption=f"📋 {summary_type_ru.title()} саммари готово!\n\n@CyberKitty19_bot"
                    )

                # Удаляем временный файл
                docx_path.unlink(missing_ok=True)
            else:
                await query.message.reply_text(summary_text, parse_mode='Markdown')

            # Обновляем сообщение с кнопками
            await query.edit_message_text(
                f"✅ {summary_type_ru.title()} саммари готово!\n\n"
                f"Что дальше будем с этим делать? 🤔"
            )

        finally:
            db.close()

    except Exception as e:
        logger.error(f"Ошибка при создании саммари: {e}")
        await query.edit_message_text(
            "❌ Произошла ошибка при создании саммари.\n\n"
            "Попробуйте позже или обратитесь в поддержку."
        )

async def handle_process_transcript_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обрабатывает кнопку 'Обработать транскрипцию'."""
    query = update.callback_query
    await query.answer()

    try:
        # Отправляем сообщение с инструкцией
        instruction_text = """🔧 **Обработка транскрипции**

Опиши, что ты хочешь получить. Например: краткое изложение, пост в канал, ТЗ обсуждаемого в транскрипции проекта — что угодно.

Если нужен конкретный формат, пришли пример: "Хочу вот так:
<пример>".

Просто отправь мне текстовое сообщение с описанием задачи! 📝"""

        await query.edit_message_text(instruction_text, parse_mode='Markdown')

        # Сохраняем состояние ожидания задачи в контексте
        context.user_data['waiting_for_task'] = True
        context.user_data['user_id'] = update.effective_user.id

    except Exception as e:
        logger.error(f"Ошибка при обработке кнопки 'Обработать': {e}")
        await query.edit_message_text("❌ Произошла ошибка. Попробуйте позже.")

async def handle_send_more_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обрабатывает кнопку 'Прислать ещё'."""
    query = update.callback_query
    await query.answer()

    try:
        # Отправляем сообщение с приглашением загрузить новый файл
        send_more_text = """📤 **Прислать ещё файл**

Отлично! Отправь мне новый видео или аудио файл для транскрипции.

Поддерживаемые форматы:
🎥 Видео: MP4, AVI, MOV, MKV, WebM и другие
🎵 Аудио: MP3, WAV, FLAC, AAC, OGG и другие
🎤 Голосовые сообщения

Максимальный размер файла: 2 ГБ
Максимальная длительность: 4 часа

Жду твой файл! 🐱"""

        await query.edit_message_text(send_more_text, parse_mode='Markdown')

    except Exception as e:
        logger.error(f"Ошибка при обработке кнопки 'Прислать ещё': {e}")
        await query.edit_message_text("❌ Произошла ошибка. Попробуйте позже.")

async def handle_main_menu_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обрабатывает кнопку 'Главное меню'."""
    query = update.callback_query
    await query.answer()

    try:
        # Возвращаемся в главное меню
        await back_to_start_callback(query, update.effective_user)

    except Exception as e:
        logger.error(f"Ошибка при обработке кнопки 'Главное меню': {e}")
        await query.edit_message_text("❌ Произошла ошибка. Попробуйте позже.")
