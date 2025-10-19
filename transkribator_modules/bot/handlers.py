"""
Обработчики сообщений для CyberKitty Transkribator
"""

import asyncio
import subprocess
import tempfile
import html
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, WebAppInfo
from telegram.ext import ContextTypes

from transkribator_modules.agent.dialog import ingest_and_prompt, handle_instruction
from transkribator_modules.audio.extractor import extract_audio_from_video, compress_audio_for_api
from transkribator_modules.beta.feature_flags import FEATURE_BETA_MODE
from transkribator_modules.beta.handlers import (
    handle_update as handle_beta_update,
    process_text as beta_process_text,
)
from transkribator_modules.config import (
    logger,
    MAX_FILE_SIZE_MB,
    VIDEOS_DIR,
    AUDIO_DIR,
    TRANSCRIPTIONS_DIR,
    BOT_TOKEN,
    AGENT_FIRST,
    MINIAPP_PUBLIC_URL,
)
from transkribator_modules.db.database import SessionLocal, UserService, log_telegram_event, log_event
from transkribator_modules.bot.commands import promo_codes_command
from transkribator_modules.transcribe.transcriber_v4 import transcribe_audio, format_transcript_with_llm, _basic_local_format
from transkribator_modules.utils.large_file_downloader import download_large_file, get_file_info


@dataclass(frozen=True)
class _YoutubeArtifacts:
    video_path: Path
    audio_path: Path
    transcript: str
    title: str
    video_id: str
    workspace: Path
    info: dict[str, Any]

def clean_html_entities(text: str) -> str:
    """Минимальная очистка текста: только удаление HTML-тегов.
    Не удаляем не-ASCII, чтобы не портить кириллицу. parse_mode=None.
    """
    if not text:
        return text
    return re.sub(r'<[^>]*>', '', text)


def _resolve_reply_target(update: Update):
    if getattr(update, "message", None):
        return update.message
    if getattr(update, "callback_query", None) and update.callback_query.message:
        return update.callback_query.message
    return None


async def _notify_free_quota_if_needed(update: Update, context: ContextTypes.DEFAULT_TYPE, user) -> None:
    message_text: str | None = None

    if getattr(user, "_was_created", False):
        message_text = (
            "🎁 В бесплатном тарифе доступны 3 видео в месяц. Используй их, чтобы попробовать все возможности."
        )
        setattr(user, "_was_created", False)
    elif getattr(user, "_usage_reset", False):
        message_text = "🔄 Лимит бесплатного тарифа обновился — снова доступны 3 бесплатные загрузки на этот месяц."
        setattr(user, "_usage_reset", False)

    if not message_text:
        return

    target = _resolve_reply_target(update)
    if target:
        await target.reply_text(message_text)
    else:
        effective_user = update.effective_user
        if effective_user:
            await context.bot.send_message(chat_id=effective_user.id, text=message_text)

# Поддерживаемые форматы
VIDEO_FORMATS = {'.mp4', '.avi', '.mov', '.mkv', '.webm', '.flv', '.wmv', '.m4v', '.3gp'}
AUDIO_FORMATS = {'.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a', '.wma', '.opus'}

_YOUTUBE_URL_RE = re.compile(
    r"(https?://(?:www\.)?(?:youtube\.com/(?:watch\?v=|shorts/)[\w-]+(?:[^\s]*)?|youtu\.be/[\w-]+(?:[^\s]*)?))",
    re.IGNORECASE,
)

def _schedule_background_task(
    context: ContextTypes.DEFAULT_TYPE,
    coro,  # type: ignore[var-annotated]
    *,
    description: str,
) -> None:
    """Запускает корутину в фоне и логирует необработанные исключения."""

    task = context.application.create_task(coro)

    def _on_done(finished_task: asyncio.Task) -> None:
        try:
            finished_task.result()
        except asyncio.CancelledError:
            logger.info(
                "Фоновая задача отменена",
                extra={"description": description},
            )
        except Exception as exc:  # noqa: BLE001 - хотим видеть стек
            logger.exception(
                "Ошибка фоновой задачи",
                extra={"description": description, "error": str(exc)},
            )

    task.add_done_callback(_on_done)

async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик команды /start"""
    welcome_text = """🎬 **CyberKitty Transkribator** 🐱

Привет! Я умею транскрибировать видео и аудио файлы любого размера!

**Что я умею:**
🎥 Обрабатывать видео до 2 ГБ
🎵 Работать с аудио файлами
📝 Создавать качественные транскрипции
🤖 Форматировать текст с помощью ИИ

**Как пользоваться:**
1. Отправьте мне видео или аудио файл
2. Подождите, пока я обработаю файл
3. Получите готовую транскрипцию!

Поддерживаемые форматы:
• Видео: MP4, AVI, MOV, MKV, WebM и другие
• Аудио: MP3, WAV, FLAC, AAC, OGG и другие

Отправьте /help для подробной помощи."""

    keyboard = InlineKeyboardMarkup([
        [
            InlineKeyboardButton(
                "🚀 Открыть веб-приложение",
                web_app=WebAppInfo(url=MINIAPP_PUBLIC_URL),
            )
        ]
    ])

    if update.message:
        await update.message.reply_text(welcome_text, parse_mode='Markdown', reply_markup=keyboard)
    else:
        await context.bot.send_message(
            chat_id=update.effective_chat.id,
            text=welcome_text,
            parse_mode='Markdown',
            reply_markup=keyboard,
        )
    try:
        log_telegram_event(
            update.effective_user,
            "command_start",
            {"chat_id": update.effective_chat.id if update.effective_chat else None},
        )
    except Exception:  # noqa: BLE001
        logger.debug("Failed to log /start event", exc_info=True)

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик команды /help"""
    help_text = """📖 **Справка по CyberKitty Transkribator**

**Основные возможности:**
• Транскрипция видео и аудио файлов
• Поддержка файлов до 2 ГБ
• Автоматическое извлечение аудио из видео
• ИИ-форматирование текста

**Команды:**
/start - Начать работу
/help - Показать эту справку
/status - Проверить статус бота

**Поддерживаемые форматы:**

🎥 **Видео:** MP4, AVI, MOV, MKV, WebM, FLV, WMV, M4V, 3GP
🎵 **Аудио:** MP3, WAV, FLAC, AAC, OGG, M4A, WMA, OPUS

**Ограничения:**
• Максимальный размер файла: 2 ГБ
• Максимальная длительность: 4 часа

**Как это работает:**
1. Вы отправляете файл
2. Если это видео - я извлекаю аудио
3. Аудио отправляется в AI API для транскрипции
4. Текст форматируется с помощью LLM
5. Вы получаете готовую транскрипцию

Просто отправьте файл и я начну обработку! 🚀"""

    await update.message.reply_text(help_text, parse_mode='Markdown')
    try:
        log_telegram_event(
            update.effective_user,
            "command_help",
            {"chat_id": update.effective_chat.id if update.effective_chat else None},
        )
    except Exception:  # noqa: BLE001
        logger.debug("Failed to log /help event", exc_info=True)

async def status_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик команды /status"""
    status_text = """✅ **Статус CyberKitty Transkribator**

🤖 Бот: Активен
🌐 Telegram Bot API Server: Активен
🎵 Обработка аудио: Доступна
🎥 Обработка видео: Доступна
🧠 ИИ транскрипция: Подключена
📝 ИИ форматирование: Активно

**Настройки:**
• Макс. размер файла: 2 ГБ
• Макс. длительность: 4 часа
• Форматы видео: 9 поддерживаемых
• Форматы аудио: 8 поддерживаемых

Готов к работе! 🚀"""

    await update.message.reply_text(status_text, parse_mode='Markdown')
    try:
        log_telegram_event(
            update.effective_user,
            "command_status",
            {"chat_id": update.effective_chat.id if update.effective_chat else None},
        )
    except Exception:  # noqa: BLE001
        logger.debug("Failed to log /status event", exc_info=True)

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик документов (файлов)"""
    document = update.message.document

    if not document:
        await update.message.reply_text("❌ Не удалось получить информацию о файле.")
        return

    # Проверяем размер файла
    file_size_mb = document.file_size / (1024 * 1024) if document.file_size else 0

    if file_size_mb > MAX_FILE_SIZE_MB:
        await update.message.reply_text(
            f"❌ Файл слишком большой: {file_size_mb:.1f} МБ\n"
            f"Максимальный размер: {MAX_FILE_SIZE_MB} МБ"
        )
        return

    # Определяем тип файла по расширению
    file_extension = Path(document.file_name).suffix.lower() if document.file_name else ''

    if file_extension in VIDEO_FORMATS:
        _schedule_background_task(
            context,
            process_video_file(update, context, document),
            description="document_video_processing",
        )
    elif file_extension in AUDIO_FORMATS:
        _schedule_background_task(
            context,
            process_audio_file(update, context, document),
            description="document_audio_processing",
        )
    else:
        await update.message.reply_text(
            f"❌ Неподдерживаемый формат файла: {file_extension}\n\n"
            f"Поддерживаемые форматы:\n"
            f"🎥 Видео: {', '.join(sorted(VIDEO_FORMATS))}\n"
            f"🎵 Аудио: {', '.join(sorted(AUDIO_FORMATS))}"
        )

async def handle_video(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик видео файлов"""
    video = update.message.video

    if not video:
        await update.message.reply_text("❌ Не удалось получить информацию о видео.")
        return

    # Проверяем размер файла
    file_size_mb = video.file_size / (1024 * 1024) if video.file_size else 0

    if file_size_mb > MAX_FILE_SIZE_MB:
        await update.message.reply_text(
            f"❌ Видео слишком большое: {file_size_mb:.1f} МБ\n"
            f"Максимальный размер: {MAX_FILE_SIZE_MB} МБ"
        )
        return

    _schedule_background_task(
        context,
        process_video_file(update, context, video),
        description="video_processing",
    )

async def handle_audio(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик аудио файлов"""
    audio = update.message.audio or update.message.voice

    if not audio:
        await update.message.reply_text("❌ Не удалось получить информацию об аудио.")
        return

    # Проверяем размер файла
    file_size_mb = audio.file_size / (1024 * 1024) if audio.file_size else 0

    if file_size_mb > MAX_FILE_SIZE_MB:
        await update.message.reply_text(
            f"❌ Аудио слишком большое: {file_size_mb:.1f} МБ\n"
            f"Максимальный размер: {MAX_FILE_SIZE_MB} МБ"
        )
        return

    _schedule_background_task(
        context,
        process_audio_file(update, context, audio),
        description="audio_processing",
    )

async def process_video_file(update: Update, context: ContextTypes.DEFAULT_TYPE, video_file, beta_enabled: bool | None = None) -> None:
    """Обрабатывает видео файл"""
    try:
        file_size_mb = video_file.file_size / (1024 * 1024) if video_file.file_size else 0
        filename = getattr(video_file, 'file_name', f"video_{video_file.file_id}")
        if beta_enabled is None:
            beta_enabled = await _is_beta_enabled(update)

        # В группах не показываем статусные сообщения
        is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")
        status_msg = None
        if not is_group:
            status_msg = await update.message.reply_text(
                f"🎬 Обрабатываю видео: {filename}\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"⏳ Скачиваю файл..."
            )

        # Создаем временные пути
        video_path = VIDEOS_DIR / f"telegram_video_{video_file.file_id}.mp4"
        audio_path = AUDIO_DIR / f"telegram_audio_{video_file.file_id}.wav"

        # Обновляем статус с информацией о скачивании
        if status_msg:
            await status_msg.edit_text(
                f"🎬 Обрабатываю видео: {filename}\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"⬇️ Скачиваю файл... (это может занять несколько минут)"
            )

        # Скачиваем файл через нашу утилиту для больших файлов
        logger.info(f"📥 Начинаю скачивание файла {filename} размером {file_size_mb:.1f} МБ")

        success = await download_large_file(
            bot_token=BOT_TOKEN,
            file_id=video_file.file_id,
            destination=video_path
        )

        if not success:
            if status_msg:
                await status_msg.edit_text("❌ Не удалось скачать файл")
            else:
                await update.message.reply_text("❌ Не удалось скачать файл")
            return

        logger.info(f"✅ Файл {filename} успешно скачан")

        # Обновляем статус
        if status_msg:
            await status_msg.edit_text(
                f"🎬 Обрабатываю видео: {filename}\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"🎵 Извлекаю аудио..."
            )

        # Извлекаем аудио
        if not await extract_audio_from_video(video_path, audio_path):
            if status_msg:
                await status_msg.edit_text("❌ Не удалось извлечь аудио из видео")
            else:
                await update.message.reply_text("❌ Не удалось извлечь аудио из видео")
            return

        # Сжимаем аудио
        if status_msg:
            await status_msg.edit_text(
                f"🎬 Обрабатываю видео: {filename}\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"🗜️ Сжимаю аудио..."
            )

        compressed_audio = await compress_audio_for_api(audio_path)

        # Транскрибируем
        if status_msg:
            await status_msg.edit_text(
                f"🎬 Обрабатываю видео: {filename}\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"🤖 Создаю транскрипцию..."
            )

        transcript = await transcribe_audio(compressed_audio)

        if not transcript or not transcript.strip():
            logger.error(f"Транскрипция не получена для видео {filename}")
            error_text = (
                "❌ Не удалось транскрибировать видео. Сервис распознавания пока недоступен. "
                "Попробуй ещё раз через пару минут."
            )
            if status_msg:
                await status_msg.edit_text(error_text)
            else:
                await update.message.reply_text(error_text)
            try:
                video_path.unlink(missing_ok=True)
                audio_path.unlink(missing_ok=True)
                comp_path = Path(compressed_audio)
                if comp_path != audio_path:
                    comp_path.unlink(missing_ok=True)
            except Exception as clear_exc:
                logger.warning(f"Не удалось удалить временные файлы после ошибки (video): {clear_exc}")
            return

        if AGENT_FIRST and transcript and transcript.strip():
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова! Открываю диалог…")
            await ingest_and_prompt(update, context, transcript, source='video')
            try:
                video_path.unlink(missing_ok=True)
                audio_path.unlink(missing_ok=True)
                comp_path = Path(compressed_audio)
                if comp_path != audio_path:
                    comp_path.unlink(missing_ok=True)
            except Exception as clear_exc:
                logger.warning(f"Не удалось удалить временные файлы (agent video): {clear_exc}")
            return
        if beta_enabled and transcript and transcript.strip() and not AGENT_FIRST:
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова! Открываю меню обработки…")
            await beta_process_text(update, context, transcript, source='video')
            try:
                video_path.unlink(missing_ok=True)
                audio_path.unlink(missing_ok=True)
                comp_path = Path(compressed_audio)
                if comp_path != audio_path:
                    comp_path.unlink(missing_ok=True)
            except Exception as clear_exc:
                logger.warning(f"Не удалось удалить временные файлы (beta video): {clear_exc}")
            return

        # Форматируем транскрипт для читаемости (OpenRouter/DeepSeek) с локальным fallback
        logger.info("Запускаю LLM-форматирование транскрипта (video)")
        formatted_transcript = None
        try:
            if transcript:
                formatted_transcript = await format_transcript_with_llm(transcript)
        except Exception as e:
            logger.warning(f"LLM-форматирование (video) исключение: {e}")
        if not formatted_transcript and transcript:
            logger.info("LLM недоступен/неверный ключ — применяю локальное форматирование")
            formatted_transcript = _basic_local_format(transcript)

        # Проверяем результат до сохранения и отправки
        if formatted_transcript and formatted_transcript.strip():
            # Сохраняем транскрипцию (уже отформатированную)
            transcript_path = TRANSCRIPTIONS_DIR / f"telegram_transcript_{video_file.file_id}.txt"
            transcript_path.write_text(formatted_transcript or "", encoding='utf-8')

            # Сохраняем в базу данных и обновляем счетчики
            try:
                from transkribator_modules.db.database import SessionLocal, UserService, TranscriptionService
                from transkribator_modules.db.database import get_media_duration

                db = SessionLocal()
                try:
                    user_service = UserService(db)
                    transcription_service = TranscriptionService(db)

                    # Получаем пользователя
                    user = user_service.get_or_create_user(
                        telegram_id=update.effective_user.id,
                        username=update.effective_user.username,
                        first_name=update.effective_user.first_name,
                        last_name=update.effective_user.last_name
                    )

                    # Проверяем лимиты использования
                    can_use, limit_message = user_service.check_usage_limit(user)
                    await _notify_free_quota_if_needed(update, context, user)
                    if not can_use:
                        if status_msg:
                            await status_msg.edit_text(f"❌ {limit_message}")
                        else:
                            await update.message.reply_text(f"❌ {limit_message}")
                        return

                    # Получаем реальную длительность аудио из видео
                    duration_minutes = get_media_duration(str(audio_path))

                    # Сохраняем транскрипцию в базу
                    transcription_service.save_transcription(
                        user=user,
                        filename=filename,
                        file_size_mb=file_size_mb,
                        audio_duration_minutes=duration_minutes,
                        raw_transcript=transcript or "",
                        formatted_transcript=formatted_transcript or "",
                        processing_time=0.0,
                        transcription_service="deepinfra",
                        formatting_service="llm" if formatted_transcript != transcript else "none"
                    )

                    # Обновляем счетчики использования
                    user_service.add_usage(user, duration_minutes)

                    logger.info(f"✅ Транскрипция сохранена для пользователя {user.telegram_id}")
                    try:
                        log_event(
                            user,
                            "video_transcription_saved",
                            {
                                "filename": filename,
                                "duration_minutes": duration_minutes,
                                "file_size_mb": file_size_mb,
                            },
                        )
                    except Exception:
                        logger.debug("Failed to log video transcription event", exc_info=True)

                finally:
                    db.close()
            except Exception as e:
                logger.error(f"Ошибка при сохранении транскрипции: {e}")

        # Проверяем, что транскрипция не пустая
        if not transcript or not transcript.strip():
            if status_msg:
                await status_msg.edit_text("❌ Не удалось создать транскрипцию")
            else:
                await update.message.reply_text("❌ Не удалось создать транскрипцию")
            return

        # Отправляем результат
        if status_msg:
            await status_msg.edit_text("✅ Транскрипция готова!")

        # Если текст короткий, отправляем в сообщении
        clean_transcript = clean_html_entities((formatted_transcript or ""))
        full_message = f"📝 Транскрипция:\n\n{clean_transcript}\n\n@CyberKitty19_bot"

        logger.info(f"Длина исходной транскрипции: {len(transcript or '')} символов")
        logger.info(f"Длина отформатированной транскрипции: {len(formatted_transcript or '')} символов")
        logger.info(f"Длина полного сообщения: {len(full_message)} символов")

        # Проверяем длину исходной транскрипции для решения о формате отправки
        if False:
            logger.info("Отправляем транскрипцию как текстовое сообщение (disabled)")
            await update.message.reply_text(full_message)
        else:
            # Если длинный, отправляем .docx
            logger.info("Отправляем транскрипцию как .docx файл")
            from docx import Document
            # Убеждаемся, что директория существует
            TRANSCRIPTIONS_DIR.mkdir(parents=True, exist_ok=True)
            docx_path = TRANSCRIPTIONS_DIR / f"transcript_{Path(filename).stem}.docx"
            document = Document()
            for line in (formatted_transcript or "").split('\n'):
                document.add_paragraph(line)
            document.save(docx_path)
            logger.info(f"Создан .docx файл: {docx_path}")
            with open(docx_path, 'rb') as f:
                await update.message.reply_document(
                    document=f,
                    filename=docx_path.name,
                    caption="📝 Транскрипция готова!\n\n@CyberKitty19_bot"
                )

        # Создаем кнопки для дальнейших действий только в личных чатах
        if not is_group:
            from telegram import InlineKeyboardButton, InlineKeyboardMarkup
            keyboard = [
                [
                    InlineKeyboardButton("🔧 Обработать", callback_data=f"process_transcript_{update.effective_user.id}"),
                    InlineKeyboardButton("📤 Прислать ещё", callback_data=f"send_more_{update.effective_user.id}")
                ],
                [
                    InlineKeyboardButton("🏠 Главное меню", callback_data=f"main_menu_{update.effective_user.id}")
                ]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)

            # Отправляем сообщение с кнопками
            await update.message.reply_text(
                "Что дальше будем с этим делать? 🤔",
                reply_markup=reply_markup
            )

        # Очищаем временные файлы
        try:
            video_path.unlink(missing_ok=True)
            audio_path.unlink(missing_ok=True)
            if compressed_audio != audio_path:
                compressed_audio.unlink(missing_ok=True)
        except Exception as e:
            logger.warning(f"Не удалось удалить временные файлы: {e}")

    except Exception as e:
        error_msg = str(e)
        logger.error(f"Ошибка при обработке видео: {e}")

        # Более информативные сообщения об ошибках
        if "timed out" in error_msg.lower() or "timeout" in error_msg.lower():
            await update.message.reply_text(
                f"⏰ Таймаут при скачивании файла\n\n"
                f"Файл размером {file_size_mb:.1f} МБ слишком долго скачивается.\n"
                f"Это может происходить из-за медленного интернета или больших размеров файла.\n\n"
                f"💡 Рекомендации:\n"
                f"• Попробуйте файл поменьше (до 100 МБ)\n"
                f"• Проверьте скорость интернета\n"
                f"• Повторите попытку через несколько минут"
            )
        else:
            await update.message.reply_text(f"❌ Ошибка при обработке видео: {error_msg}")

async def process_audio_file(update: Update, context: ContextTypes.DEFAULT_TYPE, audio_file, beta_enabled: bool | None = None) -> None:
    """Обрабатывает аудио файл"""
    try:
        file_size_mb = audio_file.file_size / (1024 * 1024) if audio_file.file_size else 0
        filename = getattr(audio_file, 'file_name', f"audio_{audio_file.file_id}")
        if beta_enabled is None:
            beta_enabled = await _is_beta_enabled(update)

        # В группах не показываем статусные сообщения
        is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")
        status_msg = None
        if not is_group:
            status_msg = await update.message.reply_text(
                f"🎵 Обрабатываю аудио: {filename}\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"⏳ Скачиваю файл..."
            )

        # Создаем временный путь
        audio_path = AUDIO_DIR / f"telegram_audio_{audio_file.file_id}.mp3"

        # Обновляем статус с информацией о скачивании
        if status_msg:
            await status_msg.edit_text(
                f"🎵 Обрабатываю аудио: {filename}\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"⬇️ Скачиваю файл... (это может занять несколько минут)"
            )

        # Скачиваем файл через нашу утилиту для больших файлов
        logger.info(f"📥 Начинаю скачивание файла {filename} размером {file_size_mb:.1f} МБ")

        success = await download_large_file(
            bot_token=BOT_TOKEN,
            file_id=audio_file.file_id,
            destination=audio_path
        )

        if not success:
            if status_msg:
                await status_msg.edit_text("❌ Не удалось скачать файл")
            else:
                await update.message.reply_text("❌ Не удалось скачать файл")
            return

        logger.info(f"✅ Файл {filename} успешно скачан")

        # Сжимаем если нужно
        if status_msg:
            await status_msg.edit_text(
                f"🎵 Обрабатываю аудио: {filename}\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"🗜️ Подготавливаю аудио..."
            )

        processed_audio = await compress_audio_for_api(audio_path)

        # Транскрибируем
        if status_msg:
            await status_msg.edit_text(
                f"🎵 Обрабатываю аудио: {filename}\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"🤖 Создаю транскрипцию..."
            )

        transcript = await transcribe_audio(processed_audio)

        if not transcript or not transcript.strip():
            logger.error(f"Транскрипция не получена для аудио {filename}")
            error_text = (
                "❌ Не удалось транскрибировать аудио. Сервис распознавания пока недоступен. "
                "Попробуй ещё раз через пару минут."
            )
            if status_msg:
                await status_msg.edit_text(error_text)
            else:
                await update.message.reply_text(error_text)
            try:
                audio_path.unlink(missing_ok=True)
                proc_path = Path(processed_audio)
                if proc_path != audio_path:
                    proc_path.unlink(missing_ok=True)
            except Exception as clear_exc:
                logger.warning(f"Не удалось удалить временные файлы после ошибки (audio): {clear_exc}")
            return

        if AGENT_FIRST and transcript and transcript.strip():
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова! Открываю диалог…")
            await ingest_and_prompt(update, context, transcript, source='audio')
            try:
                audio_path.unlink(missing_ok=True)
                proc_path = Path(processed_audio)
                if proc_path != audio_path:
                    proc_path.unlink(missing_ok=True)
            except Exception as clear_exc:
                logger.warning(f"Не удалось удалить временные файлы (agent audio): {clear_exc}")
            return
        if beta_enabled and transcript and transcript.strip() and not AGENT_FIRST:
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова! Открываю меню обработки…")
            await beta_process_text(update, context, transcript, source='audio')
            try:
                audio_path.unlink(missing_ok=True)
                proc_path = Path(processed_audio)
                if proc_path != audio_path:
                    proc_path.unlink(missing_ok=True)
            except Exception as clear_exc:
                logger.warning(f"Не удалось удалить временные файлы (beta audio): {clear_exc}")
            return

        # Форматируем транскрипт для читаемости (OpenRouter/DeepSeek) с локальным fallback
        logger.info("Запускаю LLM-форматирование транскрипта (audio)")
        formatted_transcript = None
        try:
            if transcript:
                formatted_transcript = await format_transcript_with_llm(transcript)
        except Exception as e:
            logger.warning(f"LLM-форматирование (audio) исключение: {e}")
        if not formatted_transcript and transcript:
            logger.info("LLM недоступен/неверный ключ — применяю локальное форматирование")
            formatted_transcript = _basic_local_format(transcript)

        if formatted_transcript and formatted_transcript.strip():
            # Сохраняем транскрипцию (уже отформатированную)
            transcript_path = TRANSCRIPTIONS_DIR / f"telegram_transcript_{audio_file.file_id}.txt"
            transcript_path.write_text(formatted_transcript or "", encoding='utf-8')

            # Сохраняем в базу данных и обновляем счетчики
            try:
                from transkribator_modules.db.database import SessionLocal, UserService, TranscriptionService
                from transkribator_modules.db.database import get_media_duration

                db = SessionLocal()
                try:
                    user_service = UserService(db)
                    transcription_service = TranscriptionService(db)

                    # Получаем пользователя
                    user = user_service.get_or_create_user(
                        telegram_id=update.effective_user.id,
                        username=update.effective_user.username,
                        first_name=update.effective_user.first_name,
                        last_name=update.effective_user.last_name
                    )

                    # Проверяем лимиты использования
                    can_use, limit_message = user_service.check_usage_limit(user)
                    await _notify_free_quota_if_needed(update, context, user)
                    if not can_use:
                        if status_msg:
                            await status_msg.edit_text(f"❌ {limit_message}")
                        else:
                            await update.message.reply_text(f"❌ {limit_message}")
                        return

                    # Получаем реальную длительность аудио
                    duration_minutes = get_media_duration(str(audio_path))

                    # Сохраняем транскрипцию в базу
                    transcription_service.save_transcription(
                        user=user,
                        filename=filename,
                        file_size_mb=file_size_mb,
                        audio_duration_minutes=duration_minutes,
                        raw_transcript=transcript or "",
                        formatted_transcript=formatted_transcript or "",
                        processing_time=0.0,
                        transcription_service="deepinfra",
                        formatting_service="llm" if formatted_transcript != transcript else "none"
                    )

                    # Обновляем счетчики использования
                    user_service.add_usage(user, duration_minutes)

                    logger.info(f"✅ Транскрипция сохранена для пользователя {user.telegram_id}")
                    try:
                        log_event(
                            user,
                            "audio_transcription_saved",
                            {
                                "filename": filename,
                                "duration_minutes": duration_minutes,
                                "file_size_mb": file_size_mb,
                            },
                        )
                    except Exception:
                        logger.debug("Failed to log audio transcription event", exc_info=True)

                finally:
                    db.close()
            except Exception as e:
                logger.error(f"Ошибка при сохранении транскрипции: {e}")

            # Отправляем результат
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова!")

            # Если текст короткий, отправляем в сообщении
            clean_transcript = clean_html_entities(formatted_transcript or "")
            full_message = f"📝 Транскрипция:\n\n{clean_transcript}\n\n@CyberKitty19_bot"

            logger.info(f"Длина исходной транскрипции (аудио): {len(transcript or '')} символов")
            logger.info(f"Длина отформатированной транскрипции (аудио): {len(formatted_transcript or '')} символов")
            logger.info(f"Длина полного сообщения (аудио): {len(full_message)} символов")

            # Проверяем длину исходной транскрипции для решения о формате отправки
            if False:
                logger.info("Отправляем транскрипцию как текстовое сообщение (аудио, disabled)")
                await update.message.reply_text(full_message)
            else:
                # Если длинный, отправляем .docx
                logger.info("Отправляем транскрипцию как .docx файл (аудио)")
                from docx import Document
                # Убеждаемся, что директория существует
                TRANSCRIPTIONS_DIR.mkdir(parents=True, exist_ok=True)
                docx_path = TRANSCRIPTIONS_DIR / f"transcript_{Path(filename).stem}.docx"
                document = Document()
                for line in (formatted_transcript or "").split('\n'):
                    document.add_paragraph(line)
                document.save(docx_path)
                logger.info(f"Создан .docx файл (аудио): {docx_path}")
                with open(docx_path, 'rb') as f:
                    await update.message.reply_document(
                        document=f,
                        filename=docx_path.name,
                        caption="📝 Транскрипция готова!\n\n@CyberKitty19_bot"
                    )

            # Создаем кнопки для дальнейших действий только в личных чатах
            if not is_group:
                from telegram import InlineKeyboardButton, InlineKeyboardMarkup
                keyboard = [
                    [
                        InlineKeyboardButton("🔧 Обработать", callback_data=f"process_transcript_{update.effective_user.id}"),
                        InlineKeyboardButton("📤 Прислать ещё", callback_data=f"send_more_{update.effective_user.id}")
                    ],
                    [
                        InlineKeyboardButton("🏠 Главное меню", callback_data=f"main_menu_{update.effective_user.id}")
                    ]
                ]
                reply_markup = InlineKeyboardMarkup(keyboard)

                # Отправляем сообщение с кнопками
                await update.message.reply_text(
                    "Что дальше будем с этим делать? 🤔",
                    reply_markup=reply_markup
                )
        else:
            if status_msg:
                await status_msg.edit_text("❌ Не удалось создать транскрипцию")
            else:
                await update.message.reply_text("❌ Не удалось создать транскрипцию")

        # Очищаем временные файлы
        try:
            audio_path.unlink(missing_ok=True)
            if processed_audio != audio_path:
                processed_audio.unlink(missing_ok=True)
        except Exception as e:
            logger.warning(f"Не удалось удалить временные файлы: {e}")

    except Exception as e:
        error_msg = str(e)
        logger.error(f"Ошибка при обработке аудио: {e}")

        # Более информативные сообщения об ошибках
        if "timed out" in error_msg.lower() or "timeout" in error_msg.lower():
            await update.message.reply_text(
                f"⏰ Таймаут при скачивании файла\n\n"
                f"Файл размером {file_size_mb:.1f} МБ слишком долго скачивается.\n"
                f"Это может происходить из-за медленного интернета или больших размеров файла.\n\n"
                f"💡 Рекомендации:\n"
                f"• Попробуйте файл поменьше (до 100 МБ)\n"
                f"• Проверьте скорость интернета\n"
                f"• Повторите попытку через несколько минут"
            )
        else:
            await update.message.reply_text(f"❌ Ошибка при обработке аудио: {error_msg}")

# Убрали обработчик кнопок сырой транскрипции по требованию — сосредотачиваемся на основной выдаче


def _extract_youtube_links(text: str) -> list[str]:
    if not text:
        return []
    return [match.group(1) for match in _YOUTUBE_URL_RE.finditer(text)]


async def _handle_youtube_link(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    url: str,
    beta_enabled: bool,
) -> None:
    status_msg = None
    artifacts: _YoutubeArtifacts | None = None
    is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")
    try:
        if not is_group:
            status_msg = await update.message.reply_text(
                "🎬 Нашёл ссылку на YouTube, готовлю обработку…",
                disable_web_page_preview=True,
            )
    except Exception as exc:  # noqa: BLE001
        logger.debug("Не удалось отправить статусное сообщение по YouTube", extra={"error": str(exc)})

    try:
        artifacts = await _process_youtube_ingest(update, url, status_msg)
        transcript = (artifacts.transcript or "").strip()
        summary = artifacts.title or "YouTube видео"
        filename = artifacts.video_path.name
        file_size_mb = artifacts.video_path.stat().st_size / (1024 * 1024)

        if not transcript:
            warning_text = "⚠️ Не удалось получить аудио из ролика или транскрипция пустая."
            if status_msg:
                await _safe_edit_message(status_msg, warning_text)
            else:
                await update.message.reply_text(warning_text)
            return

        if AGENT_FIRST:
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова! Открываю диалог…")
            await ingest_and_prompt(update, context, transcript, source='video')
            return

        if beta_enabled:
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова! Открываю меню обработки…")
            await beta_process_text(update, context, transcript, source='video')
            return

        logger.info("Запускаю LLM-форматирование транскрипта (youtube)")
        formatted_transcript = None
        try:
            formatted_transcript = await format_transcript_with_llm(transcript)
        except Exception as exc:  # noqa: BLE001
            logger.warning("LLM-форматирование (youtube) исключение: %s", exc)
        if not formatted_transcript:
            logger.info("LLM недоступен/неверный ключ — применяю локальное форматирование для YouTube")
            formatted_transcript = _basic_local_format(transcript)

        if formatted_transcript and formatted_transcript.strip():
            transcript_path = TRANSCRIPTIONS_DIR / f"youtube_transcript_{artifacts.video_id}.txt"
            transcript_path.write_text(formatted_transcript, encoding='utf-8')

            try:
                from transkribator_modules.db.database import (
                    SessionLocal as _SessionLocal,
                    TranscriptionService,
                    get_media_duration,
                )

                db = _SessionLocal()
                try:
                    user_service = UserService(db)
                    transcription_service = TranscriptionService(db)
                    user = user_service.get_or_create_user(
                        telegram_id=update.effective_user.id,
                        username=update.effective_user.username,
                        first_name=update.effective_user.first_name,
                        last_name=update.effective_user.last_name,
                    )

                    can_use, limit_message = user_service.check_usage_limit(user)
                    await _notify_free_quota_if_needed(update, context, user)
                    if not can_use:
                        if status_msg:
                            await status_msg.edit_text(f"❌ {limit_message}")
                        else:
                            await update.message.reply_text(f"❌ {limit_message}")
                        return

                    duration_minutes = get_media_duration(str(artifacts.audio_path))
                    transcription_service.save_transcription(
                        user=user,
                        filename=filename,
                        file_size_mb=file_size_mb,
                        audio_duration_minutes=duration_minutes,
                        raw_transcript=transcript,
                        formatted_transcript=formatted_transcript,
                        processing_time=0.0,
                        transcription_service="deepinfra",
                        formatting_service="llm" if formatted_transcript != transcript else "none",
                    )
                    user_service.add_usage(user, duration_minutes)
                    logger.info(
                        "✅ Транскрипция YouTube сохранена для пользователя %s",
                        user.telegram_id,
                    )
                finally:
                    db.close()
            except Exception as exc:  # noqa: BLE001
                logger.error("Ошибка при сохранении транскрипции YouTube: %s", exc)

            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова!")

            clean_text = clean_html_entities(formatted_transcript)
            logger.info("Длина транскрипции YouTube: %s символов", len(formatted_transcript))

            from docx import Document

            TRANSCRIPTIONS_DIR.mkdir(parents=True, exist_ok=True)
            docx_path = TRANSCRIPTIONS_DIR / f"transcript_youtube_{artifacts.video_id}.docx"
            document = Document()
            document.add_heading(summary, level=1)
            document.add_paragraph(f"Источник: {url}")
            document.add_paragraph("")
            for line in clean_text.splitlines():
                document.add_paragraph(line)
            document.save(docx_path)

            with open(docx_path, 'rb') as handle:
                await update.message.reply_document(
                    document=handle,
                    filename=docx_path.name,
                    caption="📝 Транскрипция готова!\n\n@CyberKitty19_bot",
                )

            if not is_group:
                keyboard = [
                    [
                        InlineKeyboardButton("🔧 Обработать", callback_data=f"process_transcript_{update.effective_user.id}"),
                        InlineKeyboardButton("📤 Прислать ещё", callback_data=f"send_more_{update.effective_user.id}"),
                    ],
                    [
                        InlineKeyboardButton("🏠 Главное меню", callback_data=f"main_menu_{update.effective_user.id}"),
                    ],
                ]
                await update.message.reply_text(
                    "Что дальше будем с этим делать? 🤔",
                    reply_markup=InlineKeyboardMarkup(keyboard),
                )
        else:
            message = "❌ Не удалось создать транскрипцию для YouTube видео."
            if status_msg:
                await status_msg.edit_text(message)
            else:
                await update.message.reply_text(message)
    except Exception as exc:  # noqa: BLE001
        logger.exception(
            "Не удалось обработать YouTube ссылку",
            extra={"error": str(exc), "url": url, "user_id": update.effective_user.id},
        )
        error_text = "⚠️ Не удалось обработать ссылку на YouTube. Попробуй ещё раз позже."
        if status_msg:
            await _safe_edit_message(status_msg, error_text)
        else:
            await update.message.reply_text(error_text)
    finally:
        if artifacts:
            _cleanup_workspace(artifacts.workspace)


async def _process_youtube_ingest(
    update: Update,
    url: str,
    status_msg,
) -> _YoutubeArtifacts:
    workspace = Path(tempfile.mkdtemp(prefix="youtube_ingest_"))
    try:
        await _safe_edit_message(status_msg, "📥 Скачиваю видео с YouTube…")
        download_path, info = await asyncio.to_thread(_download_youtube_media, url, workspace)

        await _safe_edit_message(status_msg, "🎛️ Конвертирую аудио…")
        wav_path = await asyncio.to_thread(_convert_to_wav, download_path)

        await _safe_edit_message(status_msg, "🗣️ Транскрибирую аудио…")
        transcript_raw = await transcribe_audio(str(wav_path))
        transcript = (transcript_raw or "").strip()
        title = (info.get("title") or "").strip() or "YouTube видео"
        video_id = info.get("id") or download_path.stem
        return _YoutubeArtifacts(
            video_path=download_path,
            audio_path=wav_path,
            transcript=transcript,
            title=title,
            video_id=video_id,
            workspace=workspace,
            info=info or {},
        )
    except Exception:
        _cleanup_workspace(workspace)
        raise


def _download_youtube_media(url: str, workspace: Path) -> tuple[Path, dict]:
    try:
        import yt_dlp  # type: ignore[import]
    except ImportError as exc:  # pragma: no cover - внешняя зависимость
        raise RuntimeError("Пакет yt-dlp не установлен для обработки ссылок YouTube.") from exc

    workspace.mkdir(parents=True, exist_ok=True)
    output_template = workspace / "%(id)s.%(ext)s"
    ydl_opts = {
        "format": "bestaudio/best",
        "outtmpl": str(output_template),
        "quiet": True,
        "no_warnings": True,
        "noplaylist": True,
    }
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=True)
        filename = ydl.prepare_filename(info)
    file_path = Path(filename)
    if not file_path.exists():
        # yt_dlp может складывать в workspace под другим расширением
        candidates = sorted(workspace.glob(f"{info.get('id', '')}.*"))
        if not candidates:
            raise FileNotFoundError("Не удалось скачать видео с YouTube")
        file_path = candidates[0]
    return file_path, info


def _convert_to_wav(input_path: Path) -> Path:
    output_path = input_path.with_suffix(".wav")
    cmd = [
        "ffmpeg",
        "-y",
        "-i",
        str(input_path),
        "-ac",
        "1",
        "-ar",
        "16000",
        str(output_path),
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    return output_path


def _cleanup_workspace(path: Path) -> None:
    for item in sorted(path.glob("**/*"), reverse=True):
        try:
            if item.is_file() or item.is_symlink():
                item.unlink()
            elif item.is_dir():
                item.rmdir()
        except FileNotFoundError:
            continue
    try:
        path.rmdir()
    except Exception:  # noqa: BLE001
        logger.debug("Не удалось полностью очистить временную директорию YouTube", exc_info=True)


async def _safe_edit_message(message, text: str) -> None:
    if not message:
        return
    try:
        await message.edit_text(text, disable_web_page_preview=True)
    except Exception as exc:  # noqa: BLE001
        logger.debug("Не удалось обновить статусное сообщение", extra={"error": str(exc)})

async def _is_beta_enabled(update: Update) -> bool:
    if not FEATURE_BETA_MODE:
        return False
    db = SessionLocal()
    try:
        user_service = UserService(db)
        db_user = user_service.get_or_create_user(
            telegram_id=update.effective_user.id,
            username=update.effective_user.username,
            first_name=update.effective_user.first_name,
            last_name=update.effective_user.last_name,
        )
        return user_service.is_beta_enabled(db_user)
    except Exception as exc:
        logger.warning(f"Не удалось проверить бета-режим: {exc}")
        return False
    finally:
        db.close()


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Интегрированный обработчик для всех типов сообщений с поддержкой Bot API Server."""
    user_id = update.effective_user.id
    chat_id = update.effective_chat.id

    # Логируем сообщение
    logger.info(f"Получено сообщение от пользователя {user_id} в чате {chat_id}")

    # Проверяем, что сообщение существует
    if not update.message:
        logger.warning("Получен update без сообщения, пропускаем")
        return

    # Определяем, работает ли бот в групповом чате
    is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")

    # Проверяем, активен ли бета-режим для пользователя (только в личных чатах)
    beta_enabled = False
    if not is_group:
        beta_enabled = await _is_beta_enabled(update)

    # Если агент ждёт инструкцию — перехватываем текстовую реплику
    if AGENT_FIRST and not is_group and update.message.text and context.user_data.get('agent_waiting_instruction'):
        await handle_instruction(update, context)
        return

    text_content = (update.message.text or update.message.caption or "").strip()

    # Команды обрабатываются телеграмом отдельно, поэтому не трогаем сообщения, начинающиеся с "@".
    if text_content.startswith("/"):
        return

    if text_content.lower().startswith("promo"):
        parts = text_content.split()
        if parts:
            # Поддерживаем как "promo CODE", так и "PROMO CODE" (без слеша).
            context.args = parts[1:]
            await promo_codes_command(update, context)
            return

    if FEATURE_BETA_MODE and beta_enabled and text_content and not AGENT_FIRST:
        youtube_links = _extract_youtube_links(text_content)
        if youtube_links:
            logger.info("Обнаружена ссылка на YouTube, запускаю бета-ингест")
            await _handle_youtube_link(update, context, youtube_links[0], beta_enabled)
            return
        logger.info("Переключение обработки в бета-режим")
        await handle_beta_update(update, context)
        return

    # Обработка видео
    if update.message.video:
        logger.info(f"Получено видео от пользователя {user_id}")
        try:
            log_telegram_event(
                update.effective_user,
                "message_video",
                {
                    "chat_id": chat_id,
                    "file_id": update.message.video.file_id,
                    "file_size": update.message.video.file_size,
                    "caption": update.message.caption,
                },
            )
        except Exception:
            logger.debug("Failed to log video message", exc_info=True)
        _schedule_background_task(
            context,
            process_video_file(update, context, update.message.video, beta_enabled=beta_enabled),
            description="message_video_processing",
        )
        return

    # Обработка аудио
    if update.message.audio:
        logger.info(f"Получено аудио от пользователя {user_id}")
        try:
            log_telegram_event(
                update.effective_user,
                "message_audio",
                {
                    "chat_id": chat_id,
                    "file_id": update.message.audio.file_id,
                    "file_size": update.message.audio.file_size,
                    "duration": update.message.audio.duration,
                    "caption": update.message.caption,
                },
            )
        except Exception:
            logger.debug("Failed to log audio message", exc_info=True)
        _schedule_background_task(
            context,
            process_audio_file(update, context, update.message.audio, beta_enabled=beta_enabled),
            description="message_audio_processing",
        )
        return

    # Обработка голосовых сообщений
    if update.message.voice:
        logger.info(f"Получено голосовое сообщение от пользователя {user_id}")
        try:
            log_telegram_event(
                update.effective_user,
                "message_voice",
                {
                    "chat_id": chat_id,
                    "file_id": update.message.voice.file_id,
                    "file_size": update.message.voice.file_size,
                    "duration": update.message.voice.duration,
                },
            )
        except Exception:
            logger.debug("Failed to log voice message", exc_info=True)
        _schedule_background_task(
            context,
            process_audio_file(update, context, update.message.voice, beta_enabled=beta_enabled),
            description="voice_processing",
        )
        return

    # Обработка документов (видео/аудио файлы)
    if update.message.document:
        document = update.message.document
        filename = document.file_name.lower() if document.file_name else ""

        # Проверяем, является ли документ видео или аудио
        if any(ext in filename for ext in VIDEO_FORMATS):
            logger.info(f"Получен видео-документ от пользователя {user_id}: {filename}")
            try:
                log_telegram_event(
                    update.effective_user,
                    "message_document_video",
                    {
                        "chat_id": chat_id,
                        "file_id": document.file_id,
                        "file_size": document.file_size,
                        "file_name": document.file_name,
                    },
                )
            except Exception:
                logger.debug("Failed to log document video", exc_info=True)
            _schedule_background_task(
                context,
                process_video_file(update, context, document, beta_enabled=beta_enabled),
                description="message_document_video_processing",
            )
            return
        elif any(ext in filename for ext in AUDIO_FORMATS):
            logger.info(f"Получен аудио-документ от пользователя {user_id}: {filename}")
            try:
                log_telegram_event(
                    update.effective_user,
                    "message_document_audio",
                    {
                        "chat_id": chat_id,
                        "file_id": document.file_id,
                        "file_size": document.file_size,
                        "file_name": document.file_name,
                    },
                )
            except Exception:
                logger.debug("Failed to log document audio", exc_info=True)
            _schedule_background_task(
                context,
                process_audio_file(update, context, document, beta_enabled=beta_enabled),
                description="message_document_audio_processing",
            )
            return

    # Если это обычное текстовое сообщение
    if update.message.text:
        # Проверяем, что это не группа или что бот упомянут в сообщении
        is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")
        bot_mentioned = False

        if is_group:
            # В группах проверяем, упомянут ли бот
            bot_username = context.bot.username
            if bot_username and f"@{bot_username}" in update.message.text:
                bot_mentioned = True

        # Отвечаем только в личных чатах или если бот упомянут в группе
        if not is_group or bot_mentioned:
            # Проверяем, ожидаем ли мы задачу для обработки транскрипции
            if context.user_data.get('waiting_for_task', False):
                await handle_transcript_processing_task(update, context)
            else:
                await update.message.reply_text(
                    "Привет! 🐱 Отправь мне видео или аудио файл, и я создам для тебя транскрипцию!\n\n"
                    "Поддерживаемые форматы:\n"
                    "📹 Видео: MP4, AVI, MOV, MKV и другие\n"
                    "🎵 Аудио: MP3, WAV, M4A, OGG и другие\n"
                    "🎤 Голосовые сообщения\n\n"
                    "Максимальный размер файла: 2 ГБ\n"
                    "Максимальная длительность: 4 часа\n\n"
                    "Используй /help для получения дополнительной информации!"
                )
            try:
                log_telegram_event(
                    update.effective_user,
                    "message_text",
                    {
                        "chat_id": chat_id,
                        "text": update.message.text,
                    },
                )
            except Exception:
                logger.debug("Failed to log text message", exc_info=True)

async def handle_transcript_processing_task(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обрабатывает текстовые сообщения с задачами для обработки транскрипции."""
    try:
        user_id = update.effective_user.id
        task_description = update.message.text

        # Сбрасываем флаг ожидания задачи
        context.user_data['waiting_for_task'] = False

        # Отправляем сообщение о начале обработки
        processing_msg = await update.message.reply_text(
            "🤖 Обрабатываю транскрипцию согласно твоей задаче...\n\n"
            "*сосредоточенно работает*\n"
            "Это может занять некоторое время...",
            parse_mode='Markdown'
        )

        # Получаем последнюю транскрипцию пользователя из базы данных
        from transkribator_modules.db.database import SessionLocal, UserService, TranscriptionService

        db = SessionLocal()
        try:
            user_service = UserService(db)
            transcription_service = TranscriptionService(db)

            # Получаем пользователя
            user = user_service.get_or_create_user(telegram_id=user_id)

            # Получаем последнюю транскрипцию пользователя
            transcriptions = transcription_service.get_user_transcriptions(user, limit=1)

            if not transcriptions:
                await processing_msg.edit_text(
                    "❌ Не найдено транскрипций для обработки.\n\n"
                    "Сначала отправьте файл для транскрипции!"
                )
                return

            latest_transcription = transcriptions[0]
            transcript_text = latest_transcription.formatted_transcript or latest_transcription.raw_transcript

            if not transcript_text:
                await processing_msg.edit_text("❌ Транскрипция пуста")
                return

            # Обрабатываем транскрипцию согласно задаче
            processed_text = await process_transcript_with_task(transcript_text, task_description)

            if not processed_text:
                await processing_msg.edit_text(
                    "❌ Не удалось обработать транскрипцию.\n\n"
                    "Возможно, сервис временно недоступен. Попробуйте позже."
                )
                return

            # Отправляем результат
            result_text = f"✅ **Результат обработки:**\n\n{processed_text}\n\n@CyberKitty19_bot"

            # Если результат длинный, отправляем файлом
            if len(result_text) > 4000:
                from docx import Document
                from pathlib import Path

                # Убеждаемся, что директория существует
                TRANSCRIPTIONS_DIR.mkdir(parents=True, exist_ok=True)
                docx_path = TRANSCRIPTIONS_DIR / f"processed_transcript_{user_id}.docx"

                document = Document()
                document.add_heading("Обработанная транскрипция", 0)
                document.add_paragraph(f"Задача: {task_description}")
                document.add_paragraph("Результат:")
                document.add_paragraph(processed_text)
                document.save(docx_path)

                with open(docx_path, 'rb') as f:
                    await update.message.reply_document(
                        document=f,
                        filename=f"processed_transcript.docx",
                        caption="✅ Результат обработки готов!\n\n@CyberKitty19_bot"
                    )

                # Удаляем временный файл
                docx_path.unlink(missing_ok=True)
            else:
                # Отправляем без parse_mode чтобы избежать ошибок с markdown entities
                await update.message.reply_text(result_text)

            # Обновляем сообщение о завершении
            await processing_msg.edit_text("✅ Обработка завершена!")

            # Показываем главное меню (личный кабинет)
            from transkribator_modules.bot.commands import personal_cabinet_command
            await personal_cabinet_command(update, context)

        finally:
            db.close()

    except Exception as e:
        logger.error(f"Ошибка при обработке задачи транскрипции: {e}")
        await update.message.reply_text(
            "❌ Произошла ошибка при обработке транскрипции.\n\n"
            "Попробуйте позже или обратитесь в поддержку."
        )

async def process_transcript_with_task(transcript_text: str, task_description: str) -> str:
    """Обрабатывает транскрипцию согласно задаче пользователя."""
    try:
        from transkribator_modules.transcribe.transcriber_v4 import request_llm_response

        system_prompt = (
            "Ты эксперт по обработке транскрипций. Твоя задача — читать запрос пользователя "
            "и выдавать готовый результат по предоставленной расшифровке.")
        user_prompt = (
            "ЗАДАЧА: {task}\n\n"
            "ТРАНСКРИПЦИЯ:\n{transcript}\n\n"
            "Обработай транскрипцию согласно задаче. Если указан конкретный формат, следуй ему точно."
        ).format(task=task_description, transcript=transcript_text)

        processed_text = None
        if request_llm_response:
            processed_text = await request_llm_response(system_prompt, user_prompt)

        if processed_text:
            cleaned_text = processed_text.strip().replace("*", "").replace("_", "").replace("`", "")
            if cleaned_text:
                return cleaned_text

        logger.warning("LLM не вернул результат для пользовательской задачи, отдаю исходную транскрипцию")
        return (
            "Не удалось обработать транскрипцию через ИИ. Вот исходная транскрипция:\n\n"
            f"{transcript_text}"
        )

    except Exception as e:
        logger.error(f"Ошибка при обработке транскрипции с задачей: {e}")
        return f"Произошла ошибка при обработке: {str(e)}"
