import asyncio
import html
import re
from pathlib import Path
from transkribator_modules.config import (
    logger, user_transcriptions, AUDIO_DIR, TRANSCRIPTIONS_DIR, MAX_MESSAGE_LENGTH, VIDEOS_DIR
)
from transkribator_modules.audio.extractor import extract_audio_from_video
from transkribator_modules.transcribe.transcriber import (
    transcribe_audio, format_transcript_with_llm, _basic_local_format
)

def clean_html_entities(text: str) -> str:
    """Минимальная очистка: убираем HTML-теги.
    Не трогаем кириллицу и другие символы, чтобы не «ломать» текст.
    При parse_mode=None Telegram не интерпретирует разметку.
    """
    if not text:
        return text
    # Удаляем только теги вида <...>
    return re.sub(r'<[^>]*>', '', text)

async def process_video_file(video_path, chat_id, message_id, context, status_message=None):
    """Обрабатывает видео из файла, извлекает аудио и выполняет транскрибацию.
    Эта версия не требует объекта Update и может быть использована напрямую с файлами."""
    
    # Кнопки убраны по требованиям — оставляем только финальную выдачу
    
    try:
        # Если на вход пришёл Telegram-объект (Video/Document), сначала скачиваем его в файл
        if not isinstance(video_path, Path):
            try:
                file_id = getattr(video_path, 'file_id', None)
                if file_id is None:
                    raise ValueError("Неверный тип аргумента video_path: ожидался Path или объект с file_id")
                # Куда скачивать видео
                download_path = VIDEOS_DIR / f"telegram_video_{message_id}.mp4"
                # Убеждаемся, что директория существует
                download_path.parent.mkdir(parents=True, exist_ok=True)
                # Скачиваем файл через Telegram Bot API
                tg_file = await context.bot.get_file(file_id)
                await tg_file.download_to_drive(custom_path=download_path)
                logger.info(f"[PROCESSOR] Видео скачано для обработки: {download_path}")
                video_path = download_path
            except Exception as download_error:
                logger.error(f"Не удалось скачать видео перед обработкой: {download_error}")
                if status_message:
                    await status_message.edit_text("Не удалось скачать видео. Попробуйте ещё раз позже.")
                else:
                    await context.bot.send_message(chat_id=chat_id, text="Не удалось скачать видео. Попробуйте ещё раз позже.")
                return
        
        # Пути к файлам
        audio_path = AUDIO_DIR / f"telegram_video_{message_id}.wav"
        
        # Проверяем наличие видео
        if not video_path.exists():
            if status_message:
                await status_message.edit_text(
                    "Мяу! Видеофайл не найден. Возможно, возникла ошибка при скачивании. *грустно машет хвостом*"
                )
            else:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Мяу! Видеофайл не найден. Возможно, возникла ошибка при скачивании. *грустно машет хвостом*"
                )
            return
        
        # Создаем статусное сообщение, если его еще нет
        if not status_message:
            status_message = await context.bot.send_message(
                chat_id=chat_id,
                text="Мур-мур! Начинаю обработку видео... *сосредоточенно смотрит на экран*"
            )
        
        # Извлекаем аудио из видео
        await status_message.edit_text(
            "Извлечение аудио из видео... *нетерпеливо перебирает лапками*"
        )
        
        audio_extracted = await extract_audio_from_video(video_path, audio_path)
        if not audio_extracted:
            await status_message.edit_text(
                "Не удалось извлечь аудио из видео. *грустно вздыхает*"
            )
            return
    
        # Транскрибируем аудио
        await status_message.edit_text(
            "Аудио извлечено! Теперь транскрибирую... *возбужденно виляет хвостом*"
        )
        
        raw_transcript = await transcribe_audio(audio_path)
        if not raw_transcript:
            await status_message.edit_text(
                "Не удалось выполнить транскрипцию аудио. *расстроенно мяукает*"
            )
            return
        
        # Форматируем транскрипцию
        await status_message.edit_text(
            "Транскрипция получена! Привожу текст в читаемый вид... *деловито стучит по клавиатуре*"
        )
        
        formatted_transcript = await format_transcript_with_llm(raw_transcript)
        if not formatted_transcript:
            # LLM недоступен или текст слишком короткий — применяем локальный форматер
            formatted_transcript = _basic_local_format(raw_transcript)
        
        # Создаем файлы с транскрипциями
        transcript_path = TRANSCRIPTIONS_DIR / f"telegram_video_{message_id}.txt"
        raw_transcript_path = TRANSCRIPTIONS_DIR / f"telegram_video_{message_id}_raw.txt"
        
        with open(transcript_path, "w", encoding="utf-8") as f:
            f.write(formatted_transcript)
            
        with open(raw_transcript_path, "w", encoding="utf-8") as f:
            f.write(raw_transcript)
        
        # Сохраняем транскрипции для пользователя
        user_transcriptions[chat_id] = {
            'raw': raw_transcript,
            'formatted': formatted_transcript,
            'path': str(transcript_path),
            'raw_path': str(raw_transcript_path),
            'timestamp': asyncio.get_event_loop().time()
        }
        
        # Отправляем результаты пользователю: короткий текст или .docx для длинного
        if len(formatted_transcript or "") > MAX_MESSAGE_LENGTH:
            if status_message:
                await status_message.edit_text(
                    "Готово! Транскрипция получилась длинной, отправляю файлом..."
                )
            from docx import Document
            docx_path = TRANSCRIPTIONS_DIR / f"transcript_video_{message_id}.docx"
            document = Document()
            for line in formatted_transcript.split('\n'):
                document.add_paragraph(line)
            document.save(docx_path)
            with open(docx_path, 'rb') as f:
                await context.bot.send_document(
                    chat_id=chat_id,
                    document=f,
                    filename=docx_path.name,
                    caption="📝 Транскрипция готова!"
                )
        else:
            clean_transcript = clean_html_entities(formatted_transcript or "")
            if status_message:
                await status_message.edit_text(
                    f"Готово! Вот транскрипция вашего видео:\n\n{clean_transcript}",
                    parse_mode=None
                )
            else:
                await context.bot.send_message(chat_id=chat_id, text=clean_transcript)
        
        logger.info(f"Транскрипция видео успешно завершена, файлы: {transcript_path}, {raw_transcript_path}")
        return transcript_path, raw_transcript_path
        
    except Exception as e:
        logger.error(f"Ошибка при обработке видео: {e}")
        import traceback
        logger.error(traceback.format_exc())
        
        if status_message:
            await status_message.edit_text(
                f"Произошла ошибка при обработке видео: {e}. *виновато опускает уши*"
            )
        return None, None

async def process_video(chat_id, message_id, update, context):
    """Обрабатывает видео, извлекает аудио и выполняет транскрибацию."""
    # Кнопки убраны — отправляем только финальный результат
    
    user_id = update.effective_user.id
    
    try:
        # Пути к файлам
        from transkribator_modules.config import VIDEOS_DIR
        video_path = VIDEOS_DIR / f"telegram_video_{message_id}.mp4"
        audio_path = AUDIO_DIR / f"telegram_video_{message_id}.wav"
        
        # Проверяем наличие видео
        if not video_path.exists():
            await update.message.reply_text(
                "Мяу! Видеофайл не найден. Возможно, возникла ошибка при скачивании. *грустно машет хвостом*"
            )
            return
        
        # Отправляем сообщение о начале обработки
        status_message = await update.message.reply_text(
            "Мур-мур! Начинаю обработку видео... *сосредоточенно смотрит на экран*"
        )
        
        # Извлекаем аудио из видео
        await status_message.edit_text(
            "Извлечение аудио из видео... *нетерпеливо перебирает лапками*"
        )
        
        audio_extracted = await extract_audio_from_video(video_path, audio_path)
        if not audio_extracted:
            await status_message.edit_text(
                "Не удалось извлечь аудио из видео. *грустно вздыхает*"
            )
            return
        
        # Транскрибируем аудио
        await status_message.edit_text(
            "Аудио извлечено! Теперь транскрибирую... *возбужденно виляет хвостом*"
        )
        
        raw_transcript = await transcribe_audio(audio_path)
        if not raw_transcript:
            await status_message.edit_text(
                "Не удалось выполнить транскрипцию аудио. *расстроенно мяукает*"
            )
            return
        
        # Форматируем транскрипцию
        await status_message.edit_text(
            "Транскрипция получена! Привожу текст в читаемый вид... *деловито стучит по клавиатуре*"
        )
        
        formatted_transcript = await format_transcript_with_llm(raw_transcript)
        
        # Создаем файлы с транскрипциями
        transcript_path = TRANSCRIPTIONS_DIR / f"telegram_video_{message_id}.txt"
        raw_transcript_path = TRANSCRIPTIONS_DIR / f"telegram_video_{message_id}_raw.txt"
        
        with open(transcript_path, "w", encoding="utf-8") as f:
            f.write(formatted_transcript)
        
        with open(raw_transcript_path, "w", encoding="utf-8") as f:
            f.write(raw_transcript)
        
        # Сохраняем транскрипции для пользователя
        user_transcriptions[user_id] = {
            'raw': raw_transcript,
            'formatted': formatted_transcript,
            'path': str(transcript_path),
            'raw_path': str(raw_transcript_path),
            'timestamp': asyncio.get_event_loop().time()
        }
        
        # Отправляем результаты пользователю: короткий текст или .docx
        if len(formatted_transcript) > MAX_MESSAGE_LENGTH:
            if status_message:
                await status_message.edit_text(
                    "Готово! Транскрипция получилась длинной, отправляю файлом..."
                )
            from docx import Document
            docx_path = TRANSCRIPTIONS_DIR / f"transcript_video_{message_id}.docx"
            document = Document()
            for line in formatted_transcript.split('\n'):
                document.add_paragraph(line)
            document.save(docx_path)
            with open(docx_path, 'rb') as f:
                await context.bot.send_document(
                    chat_id=chat_id,
                    document=f,
                    filename=docx_path.name,
                    caption="📝 Транскрипция готова!"
                )
        else:
            clean_transcript = clean_html_entities(formatted_transcript)
            if status_message:
                await status_message.edit_text(
                    f"Готово! Вот транскрипция видео:\n\n{clean_transcript}",
                    parse_mode=None
                )
            else:
                await update.message.reply_text(clean_transcript)
        
    except Exception as e:
        logger.error(f"Ошибка при обработке видео: {e}")
        await update.message.reply_text(
            f"Произошла ошибка при обработке видео: {str(e)} *испуганно прячется*"
        ) 