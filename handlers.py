"""
Обработчики сообщений для CyberKitty Transkribator
"""

import asyncio
import subprocess
import tempfile
import time
import html
import os
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import urlparse
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, WebAppInfo
from telegram.ext import ContextTypes

from transkribator_modules.agent.dialog import ingest_and_prompt, handle_instruction
from transkribator_modules.audio.extractor import extract_audio_from_video, compress_audio_for_api
from transkribator_modules.beta.feature_flags import FEATURE_BETA_MODE
from transkribator_modules.beta.handlers import (
    handle_update as handle_beta_update,
    process_text as beta_process_text,
)
from transkribator_modules.config import (
    logger,
    MAX_FILE_SIZE_MB,
    VIDEOS_DIR,
    AUDIO_DIR,
    TRANSCRIPTIONS_DIR,
    BOT_TOKEN,
    AGENT_FIRST,
    MINIAPP_PUBLIC_URL,
)
from transkribator_modules.db.database import SessionLocal, UserService, log_telegram_event, log_event
from transkribator_modules.bot.commands import promo_codes_command
from transkribator_modules.transcribe.transcriber_v4 import transcribe_audio, format_transcript_with_llm, _basic_local_format
from transkribator_modules.utils.large_file_downloader import download_large_file, get_file_info


@dataclass(frozen=True)
class _YoutubeArtifacts:
    video_path: Path
    audio_path: Path
    transcript: str
    title: str
    video_id: str
    workspace: Path
    info: dict[str, Any]

def clean_html_entities(text: str) -> str:
    """Минимальная очистка текста: только удаление HTML-тегов.
    Не удаляем не-ASCII, чтобы не портить кириллицу. parse_mode=None.
    """
    if not text:
        return text
    return re.sub(r'<[^>]*>', '', text)


def _resolve_reply_target(update: Update):
    if getattr(update, "message", None):
        return update.message
    if getattr(update, "callback_query", None) and update.callback_query.message:
        return update.callback_query.message
    return None


async def _notify_free_quota_if_needed(update: Update, context: ContextTypes.DEFAULT_TYPE, user) -> None:
    message_text: str | None = None

    if getattr(user, "_was_created", False):
        message_text = (
            "🎁 В бесплатном тарифе доступны 3 видео в месяц. Используй их, чтобы попробовать все возможности."
        )
        setattr(user, "_was_created", False)
    elif getattr(user, "_usage_reset", False):
        message_text = "🔄 Лимит бесплатного тарифа обновился — снова доступны 3 бесплатные загрузки на этот месяц."
        setattr(user, "_usage_reset", False)

    if not message_text:
        return

    target = _resolve_reply_target(update)
    if target:
        await target.reply_text(message_text)
    else:
        effective_user = update.effective_user
        if effective_user:
            await context.bot.send_message(chat_id=effective_user.id, text=message_text)

# Поддерживаемые форматы
VIDEO_FORMATS = {'.mp4', '.avi', '.mov', '.mkv', '.webm', '.flv', '.wmv', '.m4v', '.3gp'}
AUDIO_FORMATS = {'.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a', '.wma', '.opus'}

_YOUTUBE_URL_RE = re.compile(
    r"(https?://(?:www\.)?(?:youtube\.com/(?:watch\?v=|shorts/)[\w-]+(?:[^\s]*)?|youtu\.be/[\w-]+(?:[^\s]*)?))",
    re.IGNORECASE,
)
_VK_VIDEO_URL_RE = re.compile(
    r"(https?://(?:www\.)?(?:vkvideo\.ru|vk\.com)/(?:video|clip)?[-\w]+)",
    re.IGNORECASE,
)

async def generate_friendly_title_async(transcript: str, timestamp: datetime | None = None) -> str:
    """
    Асинхронно генерирует дружелюбное название транскрипции с помощью LLM.
    
    Пример: "31 окт • Встреча про проект"
    """
    import re
    
    if timestamp is None:
        timestamp = datetime.now()
    
    # Форматируем дату
    months = ['янв', 'фев', 'мар', 'апр', 'мая', 'июн', 'июл', 'авг', 'сен', 'окт', 'ноя', 'дек']
    date_str = f"{timestamp.day} {months[timestamp.month - 1]}"
    
    # Пробуем сгенерировать умное название с помощью LLM
    try:
        from transkribator_modules.transcribe.transcriber_v4 import generate_title_with_llm
        smart_title = await generate_title_with_llm(transcript)
    except Exception as e:
        logger.debug(f"Не удалось сгенерировать умное название: {e}")
        smart_title = None
    
    # Если LLM не сработал, используем простой подход - первые 2-3 слова
    if not smart_title:
        # Убираем лишние пробелы и знаки препинания
        clean_text = re.sub(r'[^\w\s\u0400-\u04FF]', ' ', transcript[:200])  # Первые 200 символов
        words = [w for w in clean_text.split() if len(w) > 2]  # Только слова длиннее 2 символов
        
        # Берём первые 2-3 значимых слова
        if len(words) >= 3:
            content = ' '.join(words[:3])
        elif len(words) >= 2:
            content = ' '.join(words[:2])
        elif len(words) >= 1:
            content = words[0]
        else:
            content = "Транскрипция"
    else:
        content = smart_title
    
    # Ограничиваем длину
    if len(content) > 40:
        content = content[:37] + "..."
    
    return f"{date_str} • {content}"

def generate_friendly_title(transcript: str, timestamp: datetime | None = None) -> str:
    """
    Синхронная обёртка для generate_friendly_title_async.
    Использует простой fallback без LLM для синхронных вызовов.
    
    Пример: "31 окт • Встреча про проект"
    """
    import re
    
    if timestamp is None:
        timestamp = datetime.now()
    
    # Форматируем дату
    months = ['янв', 'фев', 'мар', 'апр', 'мая', 'июн', 'июл', 'авг', 'сен', 'окт', 'ноя', 'дек']
    date_str = f"{timestamp.day} {months[timestamp.month - 1]}"
    
    # Простой подход - первые 2-3 слова (для синхронных вызовов)
    clean_text = re.sub(r'[^\w\s\u0400-\u04FF]', ' ', transcript[:200])
    words = [w for w in clean_text.split() if len(w) > 2]
    
    if len(words) >= 3:
        content = ' '.join(words[:3])
    elif len(words) >= 2:
        content = ' '.join(words[:2])
    elif len(words) >= 1:
        content = words[0]
    else:
        content = "Транскрипция"
    
    # Ограничиваем длину
    if len(content) > 40:
        content = content[:37] + "..."
    
    return f"{date_str} • {content}"

def _schedule_background_task(
    context: ContextTypes.DEFAULT_TYPE,
    coro,  # type: ignore[var-annotated]
    *,
    description: str,
) -> None:
    """Запускает корутину в фоне и логирует необработанные исключения."""

    task = context.application.create_task(coro)

    def _on_done(finished_task: asyncio.Task) -> None:
        try:
            finished_task.result()
        except asyncio.CancelledError:
            logger.info(
                "Фоновая задача отменена",
                extra={"description": description},
            )
        except Exception as exc:  # noqa: BLE001 - хотим видеть стек
            logger.exception(
                "Ошибка фоновой задачи",
                extra={"description": description, "error": str(exc)},
            )

    task.add_done_callback(_on_done)

async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик команды /start"""
    # Делегируем формирование приветствия модулю команд, чтобы единообразно
    # показывать обновлённое меню и уведомления про бесплатные лимиты.
    from transkribator_modules.bot.commands import start_command as commands_start_command

    await commands_start_command(update, context)
    try:
        log_telegram_event(
            update.effective_user,
            "command_start",
            {"chat_id": update.effective_chat.id if update.effective_chat else None},
        )
    except Exception:  # noqa: BLE001
        logger.debug("Failed to log /start event", exc_info=True)

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик команды /help"""
    help_text = """📖 **Справка по CyberKitty Transkribator**

**Основные возможности:**
• Транскрипция видео и аудио файлов
• Поддержка файлов до 2 ГБ
• Автоматическое извлечение аудио из видео
• ИИ-форматирование текста

**Команды:**
/start - Начать работу
/help - Показать эту справку
/status - Проверить статус бота

**Поддерживаемые форматы:**

🎥 **Видео:** MP4, AVI, MOV, MKV, WebM, FLV, WMV, M4V, 3GP
🎵 **Аудио:** MP3, WAV, FLAC, AAC, OGG, M4A, WMA, OPUS

**Ограничения:**
• Максимальный размер файла: 2 ГБ
• Максимальная длительность: 4 часа

**Как это работает:**
1. Вы отправляете файл
2. Если это видео - я извлекаю аудио
3. Аудио отправляется в AI API для транскрипции
4. Текст форматируется с помощью LLM
5. Вы получаете готовую транскрипцию

Просто отправьте файл и я начну обработку! 🚀"""

    await update.message.reply_text(help_text, parse_mode='Markdown')
    try:
        log_telegram_event(
            update.effective_user,
            "command_help",
            {"chat_id": update.effective_chat.id if update.effective_chat else None},
        )
    except Exception:  # noqa: BLE001
        logger.debug("Failed to log /help event", exc_info=True)

async def status_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик команды /status"""
    status_text = """✅ **Статус CyberKitty Transkribator**

🤖 Бот: Активен
🌐 Telegram Bot API Server: Активен
🎵 Обработка аудио: Доступна
🎥 Обработка видео: Доступна
🧠 ИИ транскрипция: Подключена
📝 ИИ форматирование: Активно

**Настройки:**
• Макс. размер файла: 2 ГБ
• Макс. длительность: 4 часа
• Форматы видео: 9 поддерживаемых
• Форматы аудио: 8 поддерживаемых

Готов к работе! 🚀"""

    await update.message.reply_text(status_text, parse_mode='Markdown')
    try:
        log_telegram_event(
            update.effective_user,
            "command_status",
            {"chat_id": update.effective_chat.id if update.effective_chat else None},
        )
    except Exception:  # noqa: BLE001
        logger.debug("Failed to log /status event", exc_info=True)

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик документов (файлов)"""
    document = update.message.document

    if not document:
        await update.message.reply_text("❌ Не удалось получить информацию о файле.")
        return

    # Логируем получение документа
    try:
        log_telegram_event(
            update.effective_user,
            "message_document",
            {
                "chat_id": update.effective_chat.id if update.effective_chat else None,
                "file_id": document.file_id,
                "file_name": document.file_name,
                "file_size": document.file_size,
                "mime_type": document.mime_type,
            }
        )
    except Exception:
        logger.debug("Failed to log document message", exc_info=True)

    # Проверяем размер файла
    file_size_mb = document.file_size / (1024 * 1024) if document.file_size else 0

    if file_size_mb > MAX_FILE_SIZE_MB:
        await update.message.reply_text(
            f"❌ Файл слишком большой: {file_size_mb:.1f} МБ\n"
            f"Максимальный размер: {MAX_FILE_SIZE_MB} МБ"
        )
        return

    # Определяем тип файла по расширению
    file_extension = Path(document.file_name).suffix.lower() if document.file_name else ''

    if file_extension in VIDEO_FORMATS:
        _schedule_background_task(
            context,
            process_video_file(update, context, document),
            description="document_video_processing",
        )
    elif file_extension in AUDIO_FORMATS:
        _schedule_background_task(
            context,
            process_audio_file(update, context, document),
            description="document_audio_processing",
        )
    else:
        await update.message.reply_text(
            f"❌ Неподдерживаемый формат файла: {file_extension}\n\n"
            f"Поддерживаемые форматы:\n"
            f"🎥 Видео: {', '.join(sorted(VIDEO_FORMATS))}\n"
            f"🎵 Аудио: {', '.join(sorted(AUDIO_FORMATS))}"
        )

async def handle_video(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик видео файлов"""
    video = update.message.video

    if not video:
        await update.message.reply_text("❌ Не удалось получить информацию о видео.")
        return

    # Проверяем размер файла
    file_size_mb = video.file_size / (1024 * 1024) if video.file_size else 0

    if file_size_mb > MAX_FILE_SIZE_MB:
        await update.message.reply_text(
            f"❌ Видео слишком большое: {file_size_mb:.1f} МБ\n"
            f"Максимальный размер: {MAX_FILE_SIZE_MB} МБ"
        )
        return

    _schedule_background_task(
        context,
        process_video_file(update, context, video),
        description="video_processing",
    )

async def handle_audio(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработчик аудио файлов"""
    audio = update.message.audio or update.message.voice

    if not audio:
        await update.message.reply_text("❌ Не удалось получить информацию об аудио.")
        return

    # Проверяем размер файла
    file_size_mb = audio.file_size / (1024 * 1024) if audio.file_size else 0

    if file_size_mb > MAX_FILE_SIZE_MB:
        await update.message.reply_text(
            f"❌ Аудио слишком большое: {file_size_mb:.1f} МБ\n"
            f"Максимальный размер: {MAX_FILE_SIZE_MB} МБ"
        )
        return

    _schedule_background_task(
        context,
        process_audio_file(update, context, audio),
        description="audio_processing",
    )

async def process_video_file(update: Update, context: ContextTypes.DEFAULT_TYPE, video_file, beta_enabled: bool | None = None) -> None:
    """Обрабатывает видео файл"""
    try:
        file_size_mb = video_file.file_size / (1024 * 1024) if video_file.file_size else 0
        filename = getattr(video_file, 'file_name', f"video_{video_file.file_id}")
        if beta_enabled is None:
            beta_enabled = await _is_beta_enabled(update)

        # В группах не показываем статусные сообщения
        is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")
        status_msg = None
        if not is_group:
            status_msg = await update.message.reply_text(
                f"🎬 Обрабатываю видео...\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"⏳ Подготовка..."
            )

        # Создаем временные пути
        video_path = VIDEOS_DIR / f"telegram_video_{video_file.file_id}.mp4"
        audio_path = AUDIO_DIR / f"telegram_audio_{video_file.file_id}.wav"

        # Обновляем статус с информацией о скачивании
        if status_msg:
            await status_msg.edit_text(
                f"🎬 Обрабатываю видео...\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"⬇️ Скачиваю файл..."
            )

        # Скачиваем файл через нашу утилиту для больших файлов
        logger.info(f"📥 Начинаю скачивание файла {filename} размером {file_size_mb:.1f} МБ")

        # Callback для отображения прогресса с throttling
        last_update = [0]  # Используем list для изменяемой переменной в closure
        
        async def progress_callback(downloaded: int, total: int):
            import time
            # Обновляем только каждые 2 секунды чтобы не спамить Telegram API
            current_time = time.time()
            if current_time - last_update[0] < 2:
                return
            last_update[0] = current_time
            
            percent = int((downloaded / total) * 100)
            # Создаём прогресс-бар из 10 квадратов
            filled = int(percent / 10)
            bar = "🟩" * filled + "⬜" * (10 - filled)
            
            if status_msg:
                try:
                    await status_msg.edit_text(
                        f"🎬 Обрабатываю видео...\n"
                        f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                        f"⬇️ Скачивание: {bar} {percent}%"
                    )
                except Exception:
                    pass  # Игнорируем ошибки редактирования (rate limit и т.д.)

        success = await download_large_file(
            bot_token=BOT_TOKEN,
            file_id=video_file.file_id,
            destination=video_path,
            progress_callback=progress_callback
        )

        if not success:
            if status_msg:
                await status_msg.edit_text("❌ Не удалось скачать файл")
            else:
                await update.message.reply_text("❌ Не удалось скачать файл")
            return

        logger.info(f"✅ Файл {filename} успешно скачан")

        # Обновляем статус
        if status_msg:
            await status_msg.edit_text(
                f"🎬 Обрабатываю видео...\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"🎵 Извлекаю аудио..."
            )

        # Извлекаем аудио
        if not await extract_audio_from_video(video_path, audio_path):
            if status_msg:
                await status_msg.edit_text("❌ Не удалось извлечь аудио из видео")
            else:
                await update.message.reply_text("❌ Не удалось извлечь аудио из видео")
            return

        # Сжимаем аудио
        if status_msg:
            await status_msg.edit_text(
                f"🎬 Обрабатываю видео...\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"🗜️ Подготавливаю аудио..."
            )

        compressed_audio = await compress_audio_for_api(audio_path)

        # Транскрибируем
        if status_msg:
            await status_msg.edit_text(
                f"🎬 Обрабатываю видео...\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"🎤 Создаю транскрипцию..."
            )

        transcript = await transcribe_audio(compressed_audio)

        if not transcript or not transcript.strip():
            logger.error(f"Транскрипция не получена для видео {filename}")
            error_text = (
                "❌ Не удалось транскрибировать видео. Сервис распознавания пока недоступен. "
                "Попробуй ещё раз через пару минут."
            )
            if status_msg:
                await status_msg.edit_text(error_text)
            else:
                await update.message.reply_text(error_text)
            try:
                video_path.unlink(missing_ok=True)
                audio_path.unlink(missing_ok=True)
                comp_path = Path(compressed_audio)
                if comp_path != audio_path:
                    comp_path.unlink(missing_ok=True)
            except Exception as clear_exc:
                logger.warning(f"Не удалось удалить временные файлы после ошибки (video): {clear_exc}")
            return

        if AGENT_FIRST and transcript and transcript.strip():
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова! Открываю диалог…")
            await ingest_and_prompt(update, context, transcript, source='video')
            try:
                video_path.unlink(missing_ok=True)
                audio_path.unlink(missing_ok=True)
                comp_path = Path(compressed_audio)
                if comp_path != audio_path:
                    comp_path.unlink(missing_ok=True)
            except Exception as clear_exc:
                logger.warning(f"Не удалось удалить временные файлы (agent video): {clear_exc}")
            return
        if beta_enabled and transcript and transcript.strip() and not AGENT_FIRST:
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова! Открываю меню обработки…")
            await beta_process_text(update, context, transcript, source='video')
            try:
                video_path.unlink(missing_ok=True)
                audio_path.unlink(missing_ok=True)
                comp_path = Path(compressed_audio)
                if comp_path != audio_path:
                    comp_path.unlink(missing_ok=True)
            except Exception as clear_exc:
                logger.warning(f"Не удалось удалить временные файлы (beta video): {clear_exc}")
            return

        # Форматируем транскрипт для читаемости (OpenRouter/DeepSeek) с локальным fallback
        logger.info("Запускаю LLM-форматирование транскрипта (video)")
        formatted_transcript = None
        try:
            if transcript:
                formatted_transcript = await format_transcript_with_llm(transcript)
        except Exception as e:
            logger.warning(f"LLM-форматирование (video) исключение: {e}")
        if not formatted_transcript and transcript:
            logger.info("LLM недоступен/неверный ключ — применяю локальное форматирование")
            formatted_transcript = _basic_local_format(transcript)

        # Проверяем результат до сохранения и отправки
        if formatted_transcript and formatted_transcript.strip():
            # Сохраняем транскрипцию (уже отформатированную)
            transcript_path = TRANSCRIPTIONS_DIR / f"telegram_transcript_{video_file.file_id}.txt"
            transcript_path.write_text(formatted_transcript or "", encoding='utf-8')

            # Сохраняем в базу данных и обновляем счетчики
            try:
                from transkribator_modules.db.database import SessionLocal, UserService, TranscriptionService
                from transkribator_modules.db.database import get_media_duration

                db = SessionLocal()
                try:
                    user_service = UserService(db)
                    transcription_service = TranscriptionService(db)

                    # Получаем пользователя
                    user = user_service.get_or_create_user(
                        telegram_id=update.effective_user.id,
                        username=update.effective_user.username,
                        first_name=update.effective_user.first_name,
                        last_name=update.effective_user.last_name
                    )

                    # Проверяем лимиты использования
                    can_use, limit_message = user_service.check_usage_limit(user)
                    await _notify_free_quota_if_needed(update, context, user)
                    if not can_use:
                        if status_msg:
                            await status_msg.edit_text(f"❌ {limit_message}")
                        else:
                            await update.message.reply_text(f"❌ {limit_message}")
                        return

                    # Получаем реальную длительность аудио из видео
                    duration_minutes = get_media_duration(str(audio_path))

                    # Сохраняем транскрипцию в базу
                    transcription_service.save_transcription(
                        user=user,
                        filename=filename,
                        file_size_mb=file_size_mb,
                        audio_duration_minutes=duration_minutes,
                        raw_transcript=transcript or "",
                        formatted_transcript=formatted_transcript or "",
                        processing_time=0.0,
                        transcription_service="deepinfra",
                        formatting_service="llm" if formatted_transcript != transcript else "none"
                    )

                    # Обновляем счетчики использования
                    user_service.add_usage(user, duration_minutes)

                    logger.info(f"✅ Транскрипция сохранена для пользователя {user.telegram_id}")
                    try:
                        log_event(
                            user,
                            "video_transcription_saved",
                            {
                                "filename": filename,
                                "duration_minutes": duration_minutes,
                                "file_size_mb": file_size_mb,
                            },
                        )
                    except Exception:
                        logger.debug("Failed to log video transcription event", exc_info=True)

                finally:
                    db.close()
            except Exception as e:
                logger.error(f"Ошибка при сохранении транскрипции: {e}")

        # Проверяем, что транскрипция не пустая
        if not transcript or not transcript.strip():
            if status_msg:
                await status_msg.edit_text("❌ Не удалось создать транскрипцию")
            else:
                await update.message.reply_text("❌ Не удалось создать транскрипцию")
            return

        # Отправляем результат
        if status_msg:
            await status_msg.edit_text("✅ Транскрипция готова!")

        # Если текст короткий, отправляем в сообщении
        clean_transcript = clean_html_entities((formatted_transcript or ""))
        full_message = f"📝 Транскрипция:\n\n{clean_transcript}\n\n@CyberKitty19_bot"

        logger.info(f"Длина исходной транскрипции: {len(transcript or '')} символов")
        logger.info(f"Длина отформатированной транскрипции: {len(formatted_transcript or '')} символов")
        logger.info(f"Длина полного сообщения: {len(full_message)} символов")

        # Проверяем длину исходной транскрипции для решения о формате отправки
        if False:
            logger.info("Отправляем транскрипцию как текстовое сообщение (disabled)")
            await update.message.reply_text(full_message)
        else:
            # Если длинный, отправляем .docx
            logger.info("Отправляем транскрипцию как .docx файл")
            from docx import Document
            from datetime import datetime
            # Убеждаемся, что директория существует
            TRANSCRIPTIONS_DIR.mkdir(parents=True, exist_ok=True)
            docx_path = TRANSCRIPTIONS_DIR / f"transcript_{Path(filename).stem}.docx"
            document = Document()
            for line in (formatted_transcript or "").split('\n'):
                document.add_paragraph(line)
            document.save(docx_path)
            logger.info(f"Создан .docx файл: {docx_path}")
            
            # Генерируем дружелюбное имя файла с помощью LLM
            friendly_name = await generate_friendly_title_async(transcript or formatted_transcript or "", datetime.now())
            friendly_filename = f"{friendly_name}.docx"
            
            with open(docx_path, 'rb') as f:
                await update.message.reply_document(
                    document=f,
                    filename=friendly_filename,
                    caption="📝 Транскрипция готова!\n\n@CyberKitty19_bot"
                )

        # Создаем кнопки для дальнейших действий только в личных чатах
        if not is_group:
            from telegram import InlineKeyboardButton, InlineKeyboardMarkup
            keyboard = [
                [
                    InlineKeyboardButton("🔧 Обработать", callback_data=f"process_transcript_{update.effective_user.id}"),
                    InlineKeyboardButton("📤 Прислать ещё", callback_data=f"send_more_{update.effective_user.id}")
                ],
                [
                    InlineKeyboardButton("🏠 Главное меню", callback_data=f"main_menu_{update.effective_user.id}")
                ]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)

            # Отправляем сообщение с кнопками
            await update.message.reply_text(
                "Что дальше будем с этим делать? 🤔",
                reply_markup=reply_markup
            )

        # Очищаем временные файлы
        try:
            video_path.unlink(missing_ok=True)
            audio_path.unlink(missing_ok=True)
            if compressed_audio != audio_path:
                compressed_audio.unlink(missing_ok=True)
        except Exception as e:
            logger.warning(f"Не удалось удалить временные файлы: {e}")

    except Exception as e:
        error_msg = str(e)
        logger.error(f"Ошибка при обработке видео: {e}")

        # Более информативные сообщения об ошибках
        if "timed out" in error_msg.lower() or "timeout" in error_msg.lower():
            await update.message.reply_text(
                f"⏰ Таймаут при скачивании файла\n\n"
                f"Файл размером {file_size_mb:.1f} МБ слишком долго скачивается.\n"
                f"Это может происходить из-за медленного интернета или больших размеров файла.\n\n"
                f"💡 Рекомендации:\n"
                f"• Попробуйте файл поменьше (до 100 МБ)\n"
                f"• Проверьте скорость интернета\n"
                f"• Повторите попытку через несколько минут"
            )
        else:
            await update.message.reply_text(f"❌ Ошибка при обработке видео: {error_msg}")

async def process_audio_file(update: Update, context: ContextTypes.DEFAULT_TYPE, audio_file, beta_enabled: bool | None = None) -> None:
    """Обрабатывает аудио файл"""
    try:
        file_size_mb = audio_file.file_size / (1024 * 1024) if audio_file.file_size else 0
        filename = getattr(audio_file, 'file_name', f"audio_{audio_file.file_id}")
        if beta_enabled is None:
            beta_enabled = await _is_beta_enabled(update)

        # В группах не показываем статусные сообщения
        is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")
        status_msg = None
        if not is_group:
            status_msg = await update.message.reply_text(
                f"🎵 Обрабатываю аудио...\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"⏳ Подготовка..."
            )

        # Создаем временный путь
        audio_path = AUDIO_DIR / f"telegram_audio_{audio_file.file_id}.mp3"

        # Обновляем статус с информацией о скачивании
        if status_msg:
            await status_msg.edit_text(
                f"🎵 Обрабатываю аудио...\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"⬇️ Скачиваю файл..."
            )

        # Скачиваем файл через нашу утилиту для больших файлов
        logger.info(f"📥 Начинаю скачивание файла {filename} размером {file_size_mb:.1f} МБ")

        # Callback для отображения прогресса с throttling
        last_update = [0]
        
        async def progress_callback(downloaded: int, total: int):
            import time
            # Обновляем только каждые 2 секунды чтобы не спамить Telegram API
            current_time = time.time()
            if current_time - last_update[0] < 2:
                return
            last_update[0] = current_time
            
            percent = int((downloaded / total) * 100)
            # Создаём прогресс-бар из 10 квадратов
            filled = int(percent / 10)
            bar = "🟩" * filled + "⬜" * (10 - filled)
            
            if status_msg:
                try:
                    await status_msg.edit_text(
                        f"🎵 Обрабатываю аудио...\n"
                        f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                        f"⬇️ Скачивание: {bar} {percent}%"
                    )
                except Exception:
                    pass  # Игнорируем ошибки редактирования (rate limit и т.д.)

        success = await download_large_file(
            bot_token=BOT_TOKEN,
            file_id=audio_file.file_id,
            destination=audio_path,
            progress_callback=progress_callback
        )

        if not success:
            if status_msg:
                await status_msg.edit_text("❌ Не удалось скачать файл")
            else:
                await update.message.reply_text("❌ Не удалось скачать файл")
            return

        logger.info(f"✅ Файл {filename} успешно скачан")

        # Сжимаем если нужно
        if status_msg:
            await status_msg.edit_text(
                f"🎵 Обрабатываю аудио...\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"🗜️ Подготавливаю аудио..."
            )

        processed_audio = await compress_audio_for_api(audio_path)

        # Транскрибируем
        if status_msg:
            await status_msg.edit_text(
                f"🎵 Обрабатываю аудио...\n"
                f"📊 Размер: {file_size_mb:.1f} МБ\n\n"
                f"🎤 Создаю транскрипцию..."
            )

        transcript = await transcribe_audio(processed_audio)

        if not transcript or not transcript.strip():
            logger.error(f"Транскрипция не получена для аудио {filename}")
            error_text = (
                "❌ Не удалось транскрибировать аудио. Сервис распознавания пока недоступен. "
                "Попробуй ещё раз через пару минут."
            )
            if status_msg:
                await status_msg.edit_text(error_text)
            else:
                await update.message.reply_text(error_text)
            try:
                audio_path.unlink(missing_ok=True)
                proc_path = Path(processed_audio)
                if proc_path != audio_path:
                    proc_path.unlink(missing_ok=True)
            except Exception as clear_exc:
                logger.warning(f"Не удалось удалить временные файлы после ошибки (audio): {clear_exc}")
            return

        if AGENT_FIRST and transcript and transcript.strip():
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова! Открываю диалог…")
            await ingest_and_prompt(update, context, transcript, source='audio')
            try:
                audio_path.unlink(missing_ok=True)
                proc_path = Path(processed_audio)
                if proc_path != audio_path:
                    proc_path.unlink(missing_ok=True)
            except Exception as clear_exc:
                logger.warning(f"Не удалось удалить временные файлы (agent audio): {clear_exc}")
            return
        if beta_enabled and transcript and transcript.strip() and not AGENT_FIRST:
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова! Открываю меню обработки…")
            await beta_process_text(update, context, transcript, source='audio')
            try:
                audio_path.unlink(missing_ok=True)
                proc_path = Path(processed_audio)
                if proc_path != audio_path:
                    proc_path.unlink(missing_ok=True)
            except Exception as clear_exc:
                logger.warning(f"Не удалось удалить временные файлы (beta audio): {clear_exc}")
            return

        # Форматируем транскрипт для читаемости (OpenRouter/DeepSeek) с локальным fallback
        logger.info("Запускаю LLM-форматирование транскрипта (audio)")
        formatted_transcript = None
        try:
            if transcript:
                formatted_transcript = await format_transcript_with_llm(transcript)
        except Exception as e:
            logger.warning(f"LLM-форматирование (audio) исключение: {e}")
        if not formatted_transcript and transcript:
            logger.info("LLM недоступен/неверный ключ — применяю локальное форматирование")
            formatted_transcript = _basic_local_format(transcript)

        if formatted_transcript and formatted_transcript.strip():
            # Сохраняем транскрипцию (уже отформатированную)
            transcript_path = TRANSCRIPTIONS_DIR / f"telegram_transcript_{audio_file.file_id}.txt"
            transcript_path.write_text(formatted_transcript or "", encoding='utf-8')

            # Сохраняем в базу данных и обновляем счетчики
            try:
                from transkribator_modules.db.database import SessionLocal, UserService, TranscriptionService
                from transkribator_modules.db.database import get_media_duration

                db = SessionLocal()
                try:
                    user_service = UserService(db)
                    transcription_service = TranscriptionService(db)

                    # Получаем пользователя
                    user = user_service.get_or_create_user(
                        telegram_id=update.effective_user.id,
                        username=update.effective_user.username,
                        first_name=update.effective_user.first_name,
                        last_name=update.effective_user.last_name
                    )

                    # Проверяем лимиты использования
                    can_use, limit_message = user_service.check_usage_limit(user)
                    await _notify_free_quota_if_needed(update, context, user)
                    if not can_use:
                        if status_msg:
                            await status_msg.edit_text(f"❌ {limit_message}")
                        else:
                            await update.message.reply_text(f"❌ {limit_message}")
                        return

                    # Получаем реальную длительность аудио
                    duration_minutes = get_media_duration(str(audio_path))

                    # Сохраняем транскрипцию в базу
                    transcription_service.save_transcription(
                        user=user,
                        filename=filename,
                        file_size_mb=file_size_mb,
                        audio_duration_minutes=duration_minutes,
                        raw_transcript=transcript or "",
                        formatted_transcript=formatted_transcript or "",
                        processing_time=0.0,
                        transcription_service="deepinfra",
                        formatting_service="llm" if formatted_transcript != transcript else "none"
                    )

                    # Обновляем счетчики использования
                    user_service.add_usage(user, duration_minutes)

                    logger.info(f"✅ Транскрипция сохранена для пользователя {user.telegram_id}")
                    try:
                        log_event(
                            user,
                            "audio_transcription_saved",
                            {
                                "filename": filename,
                                "duration_minutes": duration_minutes,
                                "file_size_mb": file_size_mb,
                            },
                        )
                    except Exception:
                        logger.debug("Failed to log audio transcription event", exc_info=True)

                finally:
                    db.close()
            except Exception as e:
                logger.error(f"Ошибка при сохранении транскрипции: {e}")

            # Отправляем результат
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова!")

            # Если текст короткий, отправляем в сообщении
            clean_transcript = clean_html_entities(formatted_transcript or "")
            full_message = f"📝 Транскрипция:\n\n{clean_transcript}\n\n@CyberKitty19_bot"

            logger.info(f"Длина исходной транскрипции (аудио): {len(transcript or '')} символов")
            logger.info(f"Длина отформатированной транскрипции (аудио): {len(formatted_transcript or '')} символов")
            logger.info(f"Длина полного сообщения (аудио): {len(full_message)} символов")

            # Проверяем длину исходной транскрипции для решения о формате отправки
            if False:
                logger.info("Отправляем транскрипцию как текстовое сообщение (аудио, disabled)")
                await update.message.reply_text(full_message)
            else:
                # Если длинный, отправляем .docx
                logger.info("Отправляем транскрипцию как .docx файл (аудио)")
                from docx import Document
                from datetime import datetime
                # Убеждаемся, что директория существует
                TRANSCRIPTIONS_DIR.mkdir(parents=True, exist_ok=True)
                docx_path = TRANSCRIPTIONS_DIR / f"transcript_{Path(filename).stem}.docx"
                document = Document()
                for line in (formatted_transcript or "").split('\n'):
                    document.add_paragraph(line)
                document.save(docx_path)
                logger.info(f"Создан .docx файл (аудио): {docx_path}")
                
                # Генерируем дружелюбное имя файла с помощью LLM
                friendly_name = await generate_friendly_title_async(transcript or formatted_transcript or "", datetime.now())
                friendly_filename = f"{friendly_name}.docx"
                
                with open(docx_path, 'rb') as f:
                    await update.message.reply_document(
                        document=f,
                        filename=friendly_filename,
                        caption="📝 Транскрипция готова!\n\n@CyberKitty19_bot"
                    )

            # Создаем кнопки для дальнейших действий только в личных чатах
            if not is_group:
                from telegram import InlineKeyboardButton, InlineKeyboardMarkup
                keyboard = [
                    [
                        InlineKeyboardButton("🔧 Обработать", callback_data=f"process_transcript_{update.effective_user.id}"),
                        InlineKeyboardButton("📤 Прислать ещё", callback_data=f"send_more_{update.effective_user.id}")
                    ],
                    [
                        InlineKeyboardButton("🏠 Главное меню", callback_data=f"main_menu_{update.effective_user.id}")
                    ]
                ]
                reply_markup = InlineKeyboardMarkup(keyboard)

                # Отправляем сообщение с кнопками
                await update.message.reply_text(
                    "Что дальше будем с этим делать? 🤔",
                    reply_markup=reply_markup
                )
        else:
            if status_msg:
                await status_msg.edit_text("❌ Не удалось создать транскрипцию")
            else:
                await update.message.reply_text("❌ Не удалось создать транскрипцию")

        # Очищаем временные файлы
        try:
            audio_path.unlink(missing_ok=True)
            if processed_audio != audio_path:
                processed_audio.unlink(missing_ok=True)
        except Exception as e:
            logger.warning(f"Не удалось удалить временные файлы: {e}")

    except Exception as e:
        error_msg = str(e)
        logger.error(f"Ошибка при обработке аудио: {e}")

        # Более информативные сообщения об ошибках
        if "timed out" in error_msg.lower() or "timeout" in error_msg.lower():
            await update.message.reply_text(
                f"⏰ Таймаут при скачивании файла\n\n"
                f"Файл размером {file_size_mb:.1f} МБ слишком долго скачивается.\n"
                f"Это может происходить из-за медленного интернета или больших размеров файла.\n\n"
                f"💡 Рекомендации:\n"
                f"• Попробуйте файл поменьше (до 100 МБ)\n"
                f"• Проверьте скорость интернета\n"
                f"• Повторите попытку через несколько минут"
            )
        else:
            await update.message.reply_text(f"❌ Ошибка при обработке аудио: {error_msg}")

# Убрали обработчик кнопок сырой транскрипции по требованию — сосредотачиваемся на основной выдаче


def _extract_youtube_links(text: str) -> list[str]:
    if not text:
        return []
    return [match.group(1) for match in _YOUTUBE_URL_RE.finditer(text)]


def _extract_vk_video_links(text: str) -> list[str]:
    if not text:
        return []
    return [match.group(1) for match in _VK_VIDEO_URL_RE.finditer(text)]


async def _handle_youtube_link(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    url: str,
    beta_enabled: bool,
) -> None:
    status_msg = None
    artifacts: _YoutubeArtifacts | None = None
    is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")
    
    # Определяем тип ссылки (YouTube или VK)
    is_vk = bool(_VK_VIDEO_URL_RE.match(url))
    platform_name = "VK видео" if is_vk else "YouTube"
    
    try:
        if not is_group:
            status_msg = await update.message.reply_text(
                f"🎬 Нашёл ссылку на {platform_name}, готовлю обработку…",
                disable_web_page_preview=True,
            )
    except Exception as exc:  # noqa: BLE001
        logger.debug(f"Не удалось отправить статусное сообщение по {platform_name}", extra={"error": str(exc)})

    try:
        artifacts = await _process_youtube_ingest(update, url, status_msg)
        transcript = (artifacts.transcript or "").strip()
        summary = artifacts.title or f"{platform_name} видео"
        filename = artifacts.video_path.name
        file_size_mb = artifacts.video_path.stat().st_size / (1024 * 1024)

        if not transcript:
            warning_text = "⚠️ Не удалось получить аудио из ролика или транскрипция пустая."
            if status_msg:
                await _safe_edit_message(status_msg, warning_text)
            else:
                await update.message.reply_text(warning_text)
            return

        if AGENT_FIRST:
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова! Открываю диалог…")
            await ingest_and_prompt(update, context, transcript, source='video')
            return

        if beta_enabled:
            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова! Открываю меню обработки…")
            await beta_process_text(update, context, transcript, source='video')
            return

        logger.info("Запускаю LLM-форматирование транскрипта (youtube)")
        formatted_transcript = None
        try:
            formatted_transcript = await format_transcript_with_llm(transcript)
        except Exception as exc:  # noqa: BLE001
            logger.warning("LLM-форматирование (youtube) исключение: %s", exc)
        if not formatted_transcript:
            logger.info("LLM недоступен/неверный ключ — применяю локальное форматирование для YouTube")
            formatted_transcript = _basic_local_format(transcript)

        if formatted_transcript and formatted_transcript.strip():
            transcript_path = TRANSCRIPTIONS_DIR / f"youtube_transcript_{artifacts.video_id}.txt"
            transcript_path.write_text(formatted_transcript, encoding='utf-8')

            try:
                from transkribator_modules.db.database import (
                    SessionLocal as _SessionLocal,
                    TranscriptionService,
                    get_media_duration,
                )

                db = _SessionLocal()
                try:
                    user_service = UserService(db)
                    transcription_service = TranscriptionService(db)
                    user = user_service.get_or_create_user(
                        telegram_id=update.effective_user.id,
                        username=update.effective_user.username,
                        first_name=update.effective_user.first_name,
                        last_name=update.effective_user.last_name,
                    )

                    can_use, limit_message = user_service.check_usage_limit(user)
                    await _notify_free_quota_if_needed(update, context, user)
                    if not can_use:
                        if status_msg:
                            await status_msg.edit_text(f"❌ {limit_message}")
                        else:
                            await update.message.reply_text(f"❌ {limit_message}")
                        return

                    duration_minutes = get_media_duration(str(artifacts.audio_path))
                    transcription_service.save_transcription(
                        user=user,
                        filename=filename,
                        file_size_mb=file_size_mb,
                        audio_duration_minutes=duration_minutes,
                        raw_transcript=transcript,
                        formatted_transcript=formatted_transcript,
                        processing_time=0.0,
                        transcription_service="deepinfra",
                        formatting_service="llm" if formatted_transcript != transcript else "none",
                    )
                    user_service.add_usage(user, duration_minutes)
                    logger.info(
                        "✅ Транскрипция YouTube сохранена для пользователя %s",
                        user.telegram_id,
                    )
                finally:
                    db.close()
            except Exception as exc:  # noqa: BLE001
                logger.error("Ошибка при сохранении транскрипции YouTube: %s", exc)

            if status_msg:
                await status_msg.edit_text("✅ Транскрипция готова!")

            clean_text = clean_html_entities(formatted_transcript)
            logger.info("Длина транскрипции YouTube: %s символов", len(formatted_transcript))

            from docx import Document
            from datetime import datetime

            TRANSCRIPTIONS_DIR.mkdir(parents=True, exist_ok=True)
            docx_path = TRANSCRIPTIONS_DIR / f"transcript_youtube_{artifacts.video_id}.docx"
            document = Document()
            document.add_heading(summary, level=1)
            document.add_paragraph(f"Источник: {url}")
            document.add_paragraph("")
            for line in clean_text.splitlines():
                document.add_paragraph(line)
            document.save(docx_path)

            # Генерируем дружелюбное имя файла для YouTube/VK с помощью LLM
            friendly_name = await generate_friendly_title_async(transcript, datetime.now())
            friendly_filename = f"{friendly_name}.docx"

            with open(docx_path, 'rb') as handle:
                await update.message.reply_document(
                    document=handle,
                    filename=friendly_filename,
                    caption="📝 Транскрипция готова!\n\n@CyberKitty19_bot",
                )

            if not is_group:
                keyboard = [
                    [
                        InlineKeyboardButton("🔧 Обработать", callback_data=f"process_transcript_{update.effective_user.id}"),
                        InlineKeyboardButton("📤 Прислать ещё", callback_data=f"send_more_{update.effective_user.id}"),
                    ],
                    [
                        InlineKeyboardButton("🏠 Главное меню", callback_data=f"main_menu_{update.effective_user.id}"),
                    ],
                ]
                await update.message.reply_text(
                    "Что дальше будем с этим делать? 🤔",
                    reply_markup=InlineKeyboardMarkup(keyboard),
                )
        else:
            message = "❌ Не удалось создать транскрипцию для YouTube видео."
            if status_msg:
                await status_msg.edit_text(message)
            else:
                await update.message.reply_text(message)
    except Exception as exc:  # noqa: BLE001
        logger.exception(
            "Не удалось обработать видео ссылку",
            extra={"error": str(exc), "url": url, "user_id": update.effective_user.id},
        )
        error_text = "⚠️ Не удалось обработать ссылку на видео. Попробуй ещё раз позже."
        if status_msg:
            await _safe_edit_message(status_msg, error_text)
        else:
            await update.message.reply_text(error_text)
    finally:
        if artifacts:
            _cleanup_workspace(artifacts.workspace)


async def _handle_gdrive_link(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    url: str,
    beta_enabled: bool,
) -> None:
    """Обрабатывает ссылку на файл Google Drive (видео/аудио)."""
    from transkribator_modules.utils.gdrive_downloader import (
        download_from_gdrive,
        GDriveDownloadError,
        extract_gdrive_id,
    )
    
    user_id = update.effective_user.id if update.effective_user else "unknown"
    file_id = extract_gdrive_id(url)
    
    logger.info(
        f"🚀 НАЧАЛО обработки Google Drive ссылки: user_id={user_id}, file_id={file_id}, beta={beta_enabled}"
    )
    start_time = time.time()
    
    status_msg = None
    workspace: Path | None = None
    is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")
    
    try:
        if not is_group:
            status_msg = await update.message.reply_text(
                "📎 Нашёл ссылку на Google Drive, начинаю скачивание…\n"
                "⏳ Это может занять несколько минут для больших файлов (2-3 GB)",
                disable_web_page_preview=True,
            )
        
        # Создаём временную директорию
        workspace = Path(tempfile.mkdtemp(prefix="gdrive_"))
        logger.info(f"📁 Создана временная директория: {workspace}")
        
        # Скачиваем файл
        await _safe_edit_message(status_msg, "📥 Скачиваю файл с Google Drive…\n⏳ Пожалуйста, подождите")
        output_path = workspace / f"gdrive_{file_id}"
        
        logger.info(f"⬇️  Начинаю скачивание в фоновом потоке: {output_path}")
        download_start = time.time()
        
        downloaded_file = await asyncio.to_thread(
            download_from_gdrive,
            url,
            output_path,
            quiet=False,
        )
        
        download_duration = time.time() - download_start
        logger.info(f"✅ Скачивание завершено за {download_duration:.1f} сек: {downloaded_file}")
        
        # Определяем тип файла
        file_ext = downloaded_file.suffix.lower()
        is_video = file_ext in ['.mp4', '.mov', '.avi', '.mkv', '.webm', '.flv', '.wmv', '.m4v']
        is_audio = file_ext in ['.mp3', '.wav', '.m4a', '.flac', '.ogg', '.opus', '.aac', '.wma']
        
        if not is_video and not is_audio:
            # Пытаемся определить по MIME type
            try:
                import magic
                mime = magic.from_file(str(downloaded_file), mime=True)
                is_video = mime.startswith('video/')
                is_audio = mime.startswith('audio/')
            except:
                pass
        
        if not is_video and not is_audio:
            error_msg = (
                f"❌ Файл имеет неподдерживаемый формат: {file_ext}\n\n"
                "Поддерживаются:\n"
                "• Видео: MP4, MOV, AVI, MKV, WebM\n"
                "• Аудио: MP3, WAV, M4A, FLAC, OGG"
            )
            await _safe_edit_message(status_msg, error_msg)
            return
        
        # Конвертируем в WAV если нужно
        if is_video or file_ext != '.wav':
            await _safe_edit_message(status_msg, "🎛️ Конвертирую аудио…")
            wav_path = await asyncio.to_thread(_convert_to_wav, downloaded_file)
        else:
            wav_path = downloaded_file
        
        # Транскрибируем
        await _safe_edit_message(status_msg, "🗣️ Транскрибирую аудио…")
        transcript_raw = await transcribe_audio(str(wav_path))
        transcript = (transcript_raw or "").strip()
        
        if not transcript:
            await _safe_edit_message(status_msg, "❌ Не удалось получить транскрипцию. Возможно, в файле нет речи.")
            return
        
        # Обрабатываем транскрипцию
        await _safe_edit_message(status_msg, "✨ Обрабатываю текст…")
        
        if beta_enabled:
            # Бета-режим: сохраняем в базу и обрабатываем через агента
            await ingest_and_prompt(
                update,
                context,
                transcript,
                source=f"gdrive:{file_id}",
            )
        else:
            # Обычный режим: генерируем саммари и отправляем
            from transkribator_modules.transcribe.postprocess import (
                summarize_text_async,
                generate_friendly_title_async,
            )
            from datetime import datetime
            
            summary = await summarize_text_async(transcript)
            clean_text = transcript.replace("  ", " ").replace("\n\n", "\n")
            
            # Создаём DOCX
            from docx import Document
            docx_path = workspace / f"transcript_gdrive_{file_id}.docx"
            document = Document()
            document.add_heading(summary or "Транскрипция", level=1)
            document.add_paragraph(f"Источник: Google Drive")
            document.add_paragraph(f"Файл: {downloaded_file.name}")
            document.add_paragraph("")
            for line in clean_text.splitlines():
                document.add_paragraph(line)
            document.save(docx_path)
            
            # Генерируем дружелюбное имя
            friendly_name = await generate_friendly_title_async(transcript, datetime.now())
            friendly_filename = f"{friendly_name}.docx"
            
            with open(docx_path, 'rb') as handle:
                await update.message.reply_document(
                    document=handle,
                    filename=friendly_filename,
                    caption=f"📄 {summary}\n\n✅ Файл успешно обработан!",
                )
        
        # Удаляем статусное сообщение
        if status_msg and not beta_enabled:
            try:
                await status_msg.delete()
            except Exception:
                pass
        
        total_duration = time.time() - start_time
        logger.info(
            f"🎉 ЗАВЕРШЕНО обработка Google Drive ссылки за {total_duration:.1f} сек: "
            f"user_id={user_id}, file_id={file_id}"
        )
    
    except GDriveDownloadError as exc:
        elapsed = time.time() - start_time
        logger.warning(
            f"⚠️ Google Drive download failed после {elapsed:.1f} сек",
            extra={"error": str(exc), "url": url, "user_id": user_id, "file_id": file_id},
        )
        error_text = str(exc)  # Ошибка уже отформатирована в GDriveDownloadError
        if status_msg:
            await _safe_edit_message(status_msg, error_text)
        else:
            await update.message.reply_text(error_text)
    
    except Exception as exc:
        elapsed = time.time() - start_time
        logger.exception(
            f"💥 Unexpected error processing Google Drive link после {elapsed:.1f} сек",
            extra={"error": str(exc), "url": url, "user_id": user_id, "file_id": file_id},
        )
        error_text = "⚠️ Не удалось обработать файл с Google Drive. Попробуйте:\n1. Отправить файл напрямую\n2. Проверить настройки доступа к файлу"
        if status_msg:
            await _safe_edit_message(status_msg, error_text)
        else:
            await update.message.reply_text(error_text)
    
    finally:
        if workspace and workspace.exists():
            _cleanup_workspace(workspace)


async def _handle_dropbox_link(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    url: str,
    beta_enabled: bool,
) -> None:
    """Обрабатывает ссылку на файл Dropbox (видео/аудио)."""
    from transkribator_modules.utils.dropbox_downloader import (
        download_from_dropbox,
        DropboxDownloadError,
        extract_dropbox_id,
    )
    
    user_id = update.effective_user.id if update.effective_user else "unknown"
    dropbox_id = extract_dropbox_id(url) or url
    
    logger.info(
        f"🚀 НАЧАЛО обработки Dropbox ссылки: user_id={user_id}, url={dropbox_id[:50]}, beta={beta_enabled}"
    )
    start_time = time.time()
    
    status_msg = None
    workspace: Path | None = None
    is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")
    
    try:
        if not is_group:
            status_msg = await update.message.reply_text(
                "📎 Нашёл ссылку на Dropbox, начинаю скачивание…\n"
                "⏳ Это может занять несколько минут для больших файлов (2-3 GB)",
                disable_web_page_preview=True,
            )
        
        # Создаём временную директорию
        workspace = Path(tempfile.mkdtemp(prefix="dropbox_"))
        logger.info(f"📁 Создана временная директория: {workspace}")
        
        # Скачиваем файл
        await _safe_edit_message(status_msg, "📥 Скачиваю файл с Dropbox…\n⏳ Пожалуйста, подождите")
        
        # Пытаемся извлечь имя файла из URL
        import urllib.parse
        parsed = urllib.parse.urlparse(url)
        path_parts = parsed.path.split('/')
        filename = path_parts[-1] if path_parts else "dropbox_file"
        output_path = workspace / filename
        
        logger.info(f"⬇️  Начинаю скачивание в фоновом потоке: {output_path}")
        download_start = time.time()
        
        downloaded_file = await asyncio.to_thread(
            download_from_dropbox,
            url,
            output_path,
        )
        
        download_duration = time.time() - download_start
        logger.info(f"✅ Скачивание завершено за {download_duration:.1f} сек: {downloaded_file}")
        
        # Определяем тип файла
        file_ext = downloaded_file.suffix.lower()
        is_video = file_ext in ['.mp4', '.mov', '.avi', '.mkv', '.webm', '.flv', '.wmv', '.m4v']
        is_audio = file_ext in ['.mp3', '.wav', '.m4a', '.flac', '.ogg', '.opus', '.aac', '.wma']
        
        if not is_video and not is_audio:
            # Пытаемся определить по MIME type
            try:
                import magic
                mime = magic.from_file(str(downloaded_file), mime=True)
                is_video = mime.startswith('video/')
                is_audio = mime.startswith('audio/')
            except:
                pass
        
        if not is_video and not is_audio:
            error_msg = (
                f"❌ Файл имеет неподдерживаемый формат: {file_ext}\n\n"
                "Поддерживаются:\n"
                "• Видео: MP4, MOV, AVI, MKV, WebM\n"
                "• Аудио: MP3, WAV, M4A, FLAC, OGG"
            )
            await _safe_edit_message(status_msg, error_msg)
            return
        
        # Конвертируем в WAV если нужно
        if is_video or file_ext != '.wav':
            await _safe_edit_message(status_msg, "🎛️ Конвертирую аудио…")
            wav_path = await asyncio.to_thread(_convert_to_wav, downloaded_file)
        else:
            wav_path = downloaded_file
        
        # Транскрибируем
        await _safe_edit_message(status_msg, "🗣️ Транскрибирую аудио…")
        transcript_raw = await transcribe_audio(str(wav_path))
        transcript = (transcript_raw or "").strip()
        
        if not transcript:
            await _safe_edit_message(status_msg, "❌ Не удалось получить транскрипцию. Возможно, в файле нет речи.")
            return
        
        # Обрабатываем транскрипцию
        await _safe_edit_message(status_msg, "✨ Обрабатываю текст…")
        
        if beta_enabled:
            # Бета-режим: сохраняем в базу и обрабатываем через агента
            await ingest_and_prompt(
                update,
                context,
                transcript,
                source=f"dropbox:{filename}",
            )
        else:
            # Обычный режим: генерируем саммари и отправляем
            from transkribator_modules.transcribe.postprocess import (
                summarize_text_async,
                generate_friendly_title_async,
            )
            from datetime import datetime
            
            summary = await summarize_text_async(transcript)
            clean_text = transcript.replace("  ", " ").replace("\n\n", "\n")
            
            # Создаём DOCX
            from docx import Document
            docx_path = workspace / f"transcript_dropbox_{filename}.docx"
            document = Document()
            document.add_heading(summary or "Транскрипция", level=1)
            document.add_paragraph(f"Источник: Dropbox")
            document.add_paragraph(f"Файл: {downloaded_file.name}")
            document.add_paragraph("")
            for line in clean_text.splitlines():
                document.add_paragraph(line)
            document.save(docx_path)
            
            # Генерируем дружелюбное имя
            friendly_name = await generate_friendly_title_async(transcript, datetime.now())
            friendly_filename = f"{friendly_name}.docx"
            
            with open(docx_path, 'rb') as handle:
                await update.message.reply_document(
                    document=handle,
                    filename=friendly_filename,
                    caption=f"📄 {summary}\n\n✅ Файл успешно обработан!",
                )
        
        # Удаляем статусное сообщение
        if status_msg and not beta_enabled:
            try:
                await status_msg.delete()
            except Exception:
                pass
        
        total_duration = time.time() - start_time
        logger.info(
            f"🎉 ЗАВЕРШЕНО обработка Dropbox ссылки за {total_duration:.1f} сек: "
            f"user_id={user_id}, filename={filename}"
        )
    
    except DropboxDownloadError as exc:
        elapsed = time.time() - start_time
        logger.warning(
            f"⚠️ Dropbox download failed после {elapsed:.1f} сек",
            extra={"error": str(exc), "url": url, "user_id": user_id},
        )
        error_text = str(exc)  # Ошибка уже отформатирована в DropboxDownloadError
        if status_msg:
            await _safe_edit_message(status_msg, error_text)
        else:
            await update.message.reply_text(error_text)
    
    except Exception as exc:
        elapsed = time.time() - start_time
        logger.exception(
            f"💥 Unexpected error processing Dropbox link после {elapsed:.1f} сек",
            extra={"error": str(exc), "url": url, "user_id": user_id},
        )
        error_text = "⚠️ Не удалось обработать файл с Dropbox. Попробуйте:\n1. Отправить файл напрямую\n2. Проверить настройки доступа к файлу"
        if status_msg:
            await _safe_edit_message(status_msg, error_text)
        else:
            await update.message.reply_text(error_text)
    
    finally:
        if workspace and workspace.exists():
            _cleanup_workspace(workspace)


async def _handle_mega_link(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    url: str,
    beta_enabled: bool,
) -> None:
    """Обрабатывает ссылку на файл Mega.nz (видео/аудио до 20 GB)."""
    from transkribator_modules.utils.mega_downloader import (
        download_from_mega,
        MegaDownloadError,
        extract_mega_id,
    )
    
    user_id = update.effective_user.id if update.effective_user else "unknown"
    mega_id = extract_mega_id(url) or url
    
    logger.info(
        f"🚀 НАЧАЛО обработки Mega.nz ссылки: user_id={user_id}, url={mega_id[:50]}, beta={beta_enabled}"
    )
    start_time = time.time()
    
    status_msg = None
    workspace: Path | None = None
    is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")
    
    try:
        if not is_group:
            status_msg = await update.message.reply_text(
                "📎 Нашёл ссылку на Mega.nz, начинаю скачивание…\n"
                "⏳ Это может занять несколько минут для больших файлов (до 20 GB)\n"
                "💡 Mega.nz: анонимная квота 10 GB/день",
                disable_web_page_preview=True,
            )
        
        # Создаём временную директорию
        workspace = Path(tempfile.mkdtemp(prefix="mega_"))
        logger.info(f"📁 Создана временная директория: {workspace}")
        
        # Скачиваем файл
        await _safe_edit_message(status_msg, "📥 Скачиваю файл с Mega.nz…\n⏳ Пожалуйста, подождите")
        
        # Mega.nz сам определяет имя файла
        output_path = workspace / "mega_file"
        
        logger.info(f"⬇️  Начинаю скачивание в фоновом потоке: {output_path}")
        download_start = time.time()
        
        downloaded_file = await asyncio.to_thread(
            download_from_mega,
            url,
            output_path,
        )
        
        download_duration = time.time() - download_start
        logger.info(f"✅ Скачивание завершено за {download_duration:.1f} сек: {downloaded_file}")
        
        # Определяем тип файла
        file_ext = downloaded_file.suffix.lower()
        is_video = file_ext in ['.mp4', '.mov', '.avi', '.mkv', '.webm', '.flv', '.wmv', '.m4v']
        is_audio = file_ext in ['.mp3', '.wav', '.m4a', '.flac', '.ogg', '.opus', '.aac', '.wma']
        
        if not is_video and not is_audio:
            # Пытаемся определить по MIME type
            try:
                import magic
                mime = magic.from_file(str(downloaded_file), mime=True)
                is_video = mime.startswith('video/')
                is_audio = mime.startswith('audio/')
            except:
                pass
        
        if not is_video and not is_audio:
            error_msg = (
                f"❌ Файл имеет неподдерживаемый формат: {file_ext}\n\n"
                "Поддерживаются:\n"
                "• Видео: MP4, MOV, AVI, MKV, WebM\n"
                "• Аудио: MP3, WAV, M4A, FLAC, OGG"
            )
            await _safe_edit_message(status_msg, error_msg)
            return
        
        # Конвертируем в WAV если нужно
        if is_video or file_ext != '.wav':
            await _safe_edit_message(status_msg, "🎛️ Конвертирую аудио…")
            wav_path = await asyncio.to_thread(_convert_to_wav, downloaded_file)
        else:
            wav_path = downloaded_file
        
        # Транскрибируем
        await _safe_edit_message(status_msg, "🗣️ Транскрибирую аудио…")
        transcript_raw = await transcribe_audio(str(wav_path))
        transcript = (transcript_raw or "").strip()
        
        if not transcript:
            await _safe_edit_message(status_msg, "❌ Не удалось получить транскрипцию. Возможно, в файле нет речи.")
            return
        
        # Обрабатываем транскрипцию
        await _safe_edit_message(status_msg, "✨ Обрабатываю текст…")
        
        if beta_enabled:
            # Бета-режим: сохраняем в базу и обрабатываем через агента
            await ingest_and_prompt(
                update,
                context,
                transcript,
                source=f"mega:{downloaded_file.name}",
            )
        else:
            # Обычный режим: генерируем саммари и отправляем
            from transkribator_modules.transcribe.postprocess import (
                summarize_text_async,
                generate_friendly_title_async,
            )
            from datetime import datetime
            
            summary = await summarize_text_async(transcript)
            clean_text = transcript.replace("  ", " ").replace("\n\n", "\n")
            
            # Создаём DOCX
            from docx import Document
            docx_path = workspace / f"transcript_mega_{downloaded_file.stem}.docx"
            document = Document()
            document.add_heading(summary or "Транскрипция", level=1)
            document.add_paragraph(f"Источник: Mega.nz")
            document.add_paragraph(f"Файл: {downloaded_file.name}")
            document.add_paragraph("")
            for line in clean_text.splitlines():
                document.add_paragraph(line)
            document.save(docx_path)
            
            # Генерируем дружелюбное имя
            friendly_name = await generate_friendly_title_async(transcript, datetime.now())
            friendly_filename = f"{friendly_name}.docx"
            
            with open(docx_path, 'rb') as handle:
                await update.message.reply_document(
                    document=handle,
                    filename=friendly_filename,
                    caption=f"📄 {summary}\n\n✅ Файл успешно обработан!",
                )
        
        # Удаляем статусное сообщение
        if status_msg and not beta_enabled:
            try:
                await status_msg.delete()
            except Exception:
                pass
        
        total_duration = time.time() - start_time
        logger.info(
            f"🎉 ЗАВЕРШЕНО обработка Mega.nz ссылки за {total_duration:.1f} сек: "
            f"user_id={user_id}, filename={downloaded_file.name}"
        )
    
    except MegaDownloadError as exc:
        elapsed = time.time() - start_time
        logger.warning(
            f"⚠️ Mega.nz download failed после {elapsed:.1f} сек",
            extra={"error": str(exc), "url": url, "user_id": user_id},
        )
        error_text = str(exc)  # Ошибка уже отформатирована в MegaDownloadError
        if status_msg:
            await _safe_edit_message(status_msg, error_text)
        else:
            await update.message.reply_text(error_text)
    
    except Exception as exc:
        elapsed = time.time() - start_time
        logger.exception(
            f"💥 Unexpected error processing Mega.nz link после {elapsed:.1f} сек",
            extra={"error": str(exc), "url": url, "user_id": user_id},
        )
        error_text = "⚠️ Не удалось обработать файл с Mega.nz. Попробуйте:\n1. Отправить файл напрямую\n2. Проверить настройки доступа к файлу"
        if status_msg:
            await _safe_edit_message(status_msg, error_text)
        else:
            await update.message.reply_text(error_text)
    
    finally:
        if workspace and workspace.exists():
            _cleanup_workspace(workspace)


async def _handle_yandex_disk_link(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    url: str,
    beta_enabled: bool,
) -> None:
    """Обрабатывает ссылку на файл Яндекс.Диска (видео/аудио до 50 GB)."""
    from transkribator_modules.utils.yandex_disk_downloader import (
        download_from_yandex_disk,
        YandexDiskDownloadError,
        extract_yandex_disk_id,
    )
    
    user_id = update.effective_user.id if update.effective_user else "unknown"
    yadisk_id = extract_yandex_disk_id(url) or url
    
    logger.info(
        f"🚀 НАЧАЛО обработки Яндекс.Диск ссылки: user_id={user_id}, url={yadisk_id[:50]}, beta={beta_enabled}"
    )
    start_time = time.time()
    
    status_msg = None
    workspace: Path | None = None
    is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")
    
    try:
        if not is_group:
            status_msg = await update.message.reply_text(
                "📎 Нашёл ссылку на Яндекс.Диск, начинаю скачивание…\n"
                "⏳ Это может занять несколько минут для больших файлов (до 50 GB)\n"
                "💡 Яндекс.Диск: квота 10 GB/день",
                disable_web_page_preview=True,
            )
        
        # Создаём временную директорию
        workspace = Path(tempfile.mkdtemp(prefix="yadisk_"))
        logger.info(f"📁 Создана временная директория: {workspace}")
        
        # Скачиваем файл
        await _safe_edit_message(status_msg, "📥 Скачиваю файл с Яндекс.Диска…\n⏳ Пожалуйста, подождите")
        
        # Пытаемся извлечь имя файла из URL или используем дефолтное
        import urllib.parse
        parsed = urllib.parse.urlparse(url)
        path_parts = parsed.path.split('/')
        filename = path_parts[-1] if path_parts and path_parts[-1] else "yadisk_file"
        output_path = workspace / filename
        
        logger.info(f"⬇️  Начинаю скачивание в фоновом потоке: {output_path}")
        download_start = time.time()
        
        downloaded_file = await asyncio.to_thread(
            download_from_yandex_disk,
            url,
            output_path,
        )
        
        download_duration = time.time() - download_start
        logger.info(f"✅ Скачивание завершено за {download_duration:.1f} сек: {downloaded_file}")
        
        # Определяем тип файла
        file_ext = downloaded_file.suffix.lower()
        is_video = file_ext in ['.mp4', '.mov', '.avi', '.mkv', '.webm', '.flv', '.wmv', '.m4v']
        is_audio = file_ext in ['.mp3', '.wav', '.m4a', '.flac', '.ogg', '.opus', '.aac', '.wma']
        
        if not is_video and not is_audio:
            # Пытаемся определить по MIME type
            try:
                import magic
                mime = magic.from_file(str(downloaded_file), mime=True)
                is_video = mime.startswith('video/')
                is_audio = mime.startswith('audio/')
            except:
                pass
        
        if not is_video and not is_audio:
            error_msg = (
                f"❌ Файл имеет неподдерживаемый формат: {file_ext}\n\n"
                "Поддерживаются:\n"
                "• Видео: MP4, MOV, AVI, MKV, WebM\n"
                "• Аудио: MP3, WAV, M4A, FLAC, OGG"
            )
            await _safe_edit_message(status_msg, error_msg)
            return
        
        # Конвертируем в WAV если нужно
        if is_video or file_ext != '.wav':
            await _safe_edit_message(status_msg, "🎛️ Конвертирую аудио…")
            wav_path = await asyncio.to_thread(_convert_to_wav, downloaded_file)
        else:
            wav_path = downloaded_file
        
        # Транскрибируем
        await _safe_edit_message(status_msg, "🗣️ Транскрибирую аудио…")
        transcript_raw = await transcribe_audio(str(wav_path))
        transcript = (transcript_raw or "").strip()
        
        if not transcript:
            await _safe_edit_message(status_msg, "❌ Не удалось получить транскрипцию. Возможно, в файле нет речи.")
            return
        
        # Обрабатываем транскрипцию
        await _safe_edit_message(status_msg, "✨ Обрабатываю текст…")
        
        if beta_enabled:
            # Бета-режим: сохраняем в базу и обрабатываем через агента
            await ingest_and_prompt(
                update,
                context,
                transcript,
                source=f"yadisk:{filename}",
            )
        else:
            # Обычный режим: генерируем саммари и отправляем
            from transkribator_modules.transcribe.postprocess import (
                summarize_text_async,
                generate_friendly_title_async,
            )
            from datetime import datetime
            
            summary = await summarize_text_async(transcript)
            clean_text = transcript.replace("  ", " ").replace("\n\n", "\n")
            
            # Создаём DOCX
            from docx import Document
            docx_path = workspace / f"transcript_yadisk_{filename}.docx"
            document = Document()
            document.add_heading(summary or "Транскрипция", level=1)
            document.add_paragraph(f"Источник: Яндекс.Диск")
            document.add_paragraph(f"Файл: {downloaded_file.name}")
            document.add_paragraph("")
            for line in clean_text.splitlines():
                document.add_paragraph(line)
            document.save(docx_path)
            
            # Генерируем дружелюбное имя
            friendly_name = await generate_friendly_title_async(transcript, datetime.now())
            friendly_filename = f"{friendly_name}.docx"
            
            with open(docx_path, 'rb') as handle:
                await update.message.reply_document(
                    document=handle,
                    filename=friendly_filename,
                    caption=f"📄 {summary}\n\n✅ Файл успешно обработан!",
                )
        
        # Удаляем статусное сообщение
        if status_msg and not beta_enabled:
            try:
                await status_msg.delete()
            except Exception:
                pass
        
        total_duration = time.time() - start_time
        logger.info(
            f"🎉 ЗАВЕРШЕНО обработка Яндекс.Диск ссылки за {total_duration:.1f} сек: "
            f"user_id={user_id}, filename={filename}"
        )
    
    except YandexDiskDownloadError as exc:
        elapsed = time.time() - start_time
        logger.warning(
            f"⚠️ Yandex.Disk download failed после {elapsed:.1f} сек",
            extra={"error": str(exc), "url": url, "user_id": user_id},
        )
        error_text = str(exc)  # Ошибка уже отформатирована в YandexDiskDownloadError
        if status_msg:
            await _safe_edit_message(status_msg, error_text)
        else:
            await update.message.reply_text(error_text)
    
    except Exception as exc:
        elapsed = time.time() - start_time
        logger.exception(
            f"💥 Unexpected error processing Yandex.Disk link после {elapsed:.1f} сек",
            extra={"error": str(exc), "url": url, "user_id": user_id},
        )
        error_text = "⚠️ Не удалось обработать файл с Яндекс.Диска. Попробуйте:\n1. Отправить файл напрямую\n2. Проверить настройки доступа к файлу"
        if status_msg:
            await _safe_edit_message(status_msg, error_text)
        else:
            await update.message.reply_text(error_text)
    
    finally:
        if workspace and workspace.exists():
            _cleanup_workspace(workspace)


async def _process_youtube_ingest(
    update: Update,
    url: str,
    status_msg,
) -> _YoutubeArtifacts:
    workspace = Path(tempfile.mkdtemp(prefix="video_ingest_"))
    try:
        await _safe_edit_message(status_msg, "📥 Скачиваю видео…")
        download_path, info = await asyncio.to_thread(_download_youtube_media, url, workspace)

        await _safe_edit_message(status_msg, "🎛️ Конвертирую аудио…")
        wav_path = await asyncio.to_thread(_convert_to_wav, download_path)

        await _safe_edit_message(status_msg, "🗣️ Транскрибирую аудио…")
        transcript_raw = await transcribe_audio(str(wav_path))
        transcript = (transcript_raw or "").strip()
        title = (info.get("title") or "").strip() or "Видео"
        video_id = info.get("id") or download_path.stem
        return _YoutubeArtifacts(
            video_path=download_path,
            audio_path=wav_path,
            transcript=transcript,
            title=title,
            video_id=video_id,
            workspace=workspace,
            info=info or {},
        )
    except Exception:
        _cleanup_workspace(workspace)
        raise


def _download_youtube_media(url: str, workspace: Path) -> tuple[Path, dict]:
    try:
        import yt_dlp  # type: ignore[import]
    except ImportError as exc:  # pragma: no cover - внешняя зависимость
        raise RuntimeError("Пакет yt-dlp не установлен для обработки ссылок YouTube/VK.") from exc

    workspace.mkdir(parents=True, exist_ok=True)
    output_template = workspace / "%(id)s.%(ext)s"
    ydl_opts = {
        "format": "bestaudio/best",
        "outtmpl": str(output_template),
        "quiet": True,
        "no_warnings": True,
        "noplaylist": True,
        "user_agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    }
    proxy_url = os.getenv("YTDLP_PROXY", "").strip()
    if proxy_url:
        ydl_opts["proxy"] = proxy_url
        try:
            parsed_proxy = urlparse(proxy_url)
            logger.info(
                "🛡️ Использую прокси для YouTube загрузки",
                extra={
                    "proxy_scheme": parsed_proxy.scheme or "",
                    "proxy_host": parsed_proxy.hostname or "",
                    "proxy_port": parsed_proxy.port,
                },
            )
        except Exception:  # noqa: BLE001
            logger.info("🛡️ Использую прокси для YouTube загрузки", extra={"proxy": "custom"})
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=True)
        filename = ydl.prepare_filename(info)
    file_path = Path(filename)
    if not file_path.exists():
        # yt_dlp может складывать в workspace под другим расширением
        candidates = sorted(workspace.glob(f"{info.get('id', '')}.*"))
        if not candidates:
            raise FileNotFoundError("Не удалось скачать видео с YouTube/VK")
        file_path = candidates[0]
    return file_path, info


def _convert_to_wav(input_path: Path) -> Path:
    output_path = input_path.with_suffix(".wav")
    cmd = [
        "ffmpeg",
        "-y",
        "-i",
        str(input_path),
        "-ac",
        "1",
        "-ar",
        "16000",
        str(output_path),
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    return output_path


def _cleanup_workspace(path: Path) -> None:
    for item in sorted(path.glob("**/*"), reverse=True):
        try:
            if item.is_file() or item.is_symlink():
                item.unlink()
            elif item.is_dir():
                item.rmdir()
        except FileNotFoundError:
            continue
    try:
        path.rmdir()
    except Exception:  # noqa: BLE001
        logger.debug("Не удалось полностью очистить временную директорию YouTube", exc_info=True)


async def _safe_edit_message(message, text: str) -> None:
    if not message:
        return
    try:
        await message.edit_text(text, disable_web_page_preview=True)
    except Exception as exc:  # noqa: BLE001
        logger.debug("Не удалось обновить статусное сообщение", extra={"error": str(exc)})

async def _is_beta_enabled(update: Update) -> bool:
    if not FEATURE_BETA_MODE:
        return False
    db = SessionLocal()
    try:
        user_service = UserService(db)
        db_user = user_service.get_or_create_user(
            telegram_id=update.effective_user.id,
            username=update.effective_user.username,
            first_name=update.effective_user.first_name,
            last_name=update.effective_user.last_name,
        )
        return user_service.is_beta_enabled(db_user)
    except Exception as exc:
        logger.warning(f"Не удалось проверить бета-режим: {exc}")
        return False
    finally:
        db.close()


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Интегрированный обработчик для всех типов сообщений с поддержкой Bot API Server."""
    user_id = update.effective_user.id
    chat_id = update.effective_chat.id

    # Логируем сообщение
    logger.info(f"Получено сообщение от пользователя {user_id} в чате {chat_id}")

    # Проверяем, что сообщение существует
    if not update.message:
        logger.warning("Получен update без сообщения, пропускаем")
        return

    # Определяем, работает ли бот в групповом чате
    is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")

    # Проверяем, активен ли бета-режим для пользователя (только в личных чатах)
    beta_enabled = False
    if not is_group:
        beta_enabled = await _is_beta_enabled(update)

    # Если агент ждёт инструкцию — перехватываем текстовую реплику
    if AGENT_FIRST and not is_group and update.message.text and context.user_data.get('agent_waiting_instruction'):
        await handle_instruction(update, context)
        return

    text_content = (update.message.text or update.message.caption or "").strip()

    # Команды обрабатываются телеграмом отдельно, поэтому не трогаем сообщения, начинающиеся с "@".
    if text_content.startswith("/"):
        return

    if text_content.lower().startswith("promo"):
        parts = text_content.split()
        if parts:
            # Поддерживаем как "promo CODE", так и "PROMO CODE" (без слеша).
            context.args = parts[1:]
            await promo_codes_command(update, context)
            return

    # Проверяем облачные хранилища ПЕРЕД бета-режимом (чтобы работало всегда)
    if text_content:
        # Проверяем Google Drive ссылки
        from transkribator_modules.utils.gdrive_downloader import extract_gdrive_links
        gdrive_links = extract_gdrive_links(text_content)
        if gdrive_links:
            logger.info("Обнаружена ссылка на Google Drive, запускаю обработку")
            # Логируем Google Drive ссылку
            try:
                log_event(
                    update.effective_user.id,
                    "bot_media_gdrive_link",
                    {
                        "url": gdrive_links[0],
                        "chat_id": chat_id,
                    }
                )
            except Exception:
                logger.debug("Failed to log gdrive link event", exc_info=True)
            
            _schedule_background_task(
                context,
                _handle_gdrive_link(update, context, gdrive_links[0], beta_enabled),
                description="gdrive_link_processing",
            )
            return
        
        # Проверяем Dropbox ссылки
        from transkribator_modules.utils.dropbox_downloader import extract_dropbox_links
        dropbox_links = extract_dropbox_links(text_content)
        if dropbox_links:
            logger.info("Обнаружена ссылка на Dropbox, запускаю обработку")
            # Логируем Dropbox ссылку
            try:
                log_event(
                    update.effective_user.id,
                    "bot_media_dropbox_link",
                    {
                        "url": dropbox_links[0],
                        "chat_id": chat_id,
                    }
                )
            except Exception:
                logger.debug("Failed to log dropbox link event", exc_info=True)
            
            _schedule_background_task(
                context,
                _handle_dropbox_link(update, context, dropbox_links[0], beta_enabled),
                description="dropbox_link_processing",
            )
            return
        
        # Проверяем Mega.nz ссылки
        from transkribator_modules.utils.mega_downloader import extract_mega_links
        mega_links = extract_mega_links(text_content)
        if mega_links:
            logger.info("Обнаружена ссылка на Mega.nz, запускаю обработку")
            # Логируем Mega.nz ссылку
            try:
                log_event(
                    update.effective_user.id,
                    "bot_media_mega_link",
                    {
                        "url": mega_links[0],
                        "chat_id": chat_id,
                    }
                )
            except Exception:
                logger.debug("Failed to log mega link event", exc_info=True)
            
            _schedule_background_task(
                context,
                _handle_mega_link(update, context, mega_links[0], beta_enabled),
                description="mega_link_processing",
            )
            return
        
        # Проверяем Яндекс.Диск ссылки
        from transkribator_modules.utils.yandex_disk_downloader import extract_yandex_disk_links
        yadisk_links = extract_yandex_disk_links(text_content)
        if yadisk_links:
            logger.info("Обнаружена ссылка на Яндекс.Диск, запускаю обработку")
            # Логируем Яндекс.Диск ссылку
            try:
                log_event(
                    update.effective_user.id,
                    "bot_media_yandex_disk_link",
                    {
                        "url": yadisk_links[0],
                        "chat_id": chat_id,
                    }
                )
            except Exception:
                logger.debug("Failed to log yandex disk link event", exc_info=True)
            
            _schedule_background_task(
                context,
                _handle_yandex_disk_link(update, context, yadisk_links[0], beta_enabled),
                description="yandex_disk_link_processing",
            )
            return

    if FEATURE_BETA_MODE and beta_enabled and text_content and not AGENT_FIRST:
        youtube_links = _extract_youtube_links(text_content)
        if youtube_links:
            logger.info("Обнаружена ссылка на YouTube, запускаю бета-ингест")
            # Логируем YouTube ссылку
            try:
                log_event(
                    update.effective_user.id,
                    "bot_media_youtube_link",
                    {
                        "url": youtube_links[0],
                        "chat_id": chat_id,
                    }
                )
            except Exception:
                logger.debug("Failed to log youtube link event", exc_info=True)
            
            await _handle_youtube_link(update, context, youtube_links[0], beta_enabled)
            return
        vk_video_links = _extract_vk_video_links(text_content)
        if vk_video_links:
            logger.info("Обнаружена ссылка на VK видео, запускаю бета-ингест")
            # Логируем VK ссылку
            try:
                log_event(
                    update.effective_user.id,
                    "bot_media_vk_link",
                    {
                        "url": vk_video_links[0],
                        "chat_id": chat_id,
                    }
                )
            except Exception:
                logger.debug("Failed to log vk link event", exc_info=True)
            
            await _handle_youtube_link(update, context, vk_video_links[0], beta_enabled)
            return
        
        logger.info("Переключение обработки в бета-режим")
        await handle_beta_update(update, context)
        return

    # Обработка видео
    if update.message.video:
        logger.info(f"Получено видео от пользователя {user_id}")
        try:
            log_telegram_event(
                update.effective_user,
                "message_video",
                {
                    "chat_id": chat_id,
                    "file_id": update.message.video.file_id,
                    "file_size": update.message.video.file_size,
                    "caption": update.message.caption,
                },
            )
        except Exception:
            logger.debug("Failed to log video message", exc_info=True)
        _schedule_background_task(
            context,
            process_video_file(update, context, update.message.video, beta_enabled=beta_enabled),
            description="message_video_processing",
        )
        return

    # Обработка аудио
    if update.message.audio:
        logger.info(f"Получено аудио от пользователя {user_id}")
        try:
            log_telegram_event(
                update.effective_user,
                "message_audio",
                {
                    "chat_id": chat_id,
                    "file_id": update.message.audio.file_id,
                    "file_size": update.message.audio.file_size,
                    "duration": update.message.audio.duration,
                    "caption": update.message.caption,
                },
            )
        except Exception:
            logger.debug("Failed to log audio message", exc_info=True)
        _schedule_background_task(
            context,
            process_audio_file(update, context, update.message.audio, beta_enabled=beta_enabled),
            description="message_audio_processing",
        )
        return

    # Обработка голосовых сообщений
    if update.message.voice:
        logger.info(f"Получено голосовое сообщение от пользователя {user_id}")
        try:
            log_telegram_event(
                update.effective_user,
                "message_voice",
                {
                    "chat_id": chat_id,
                    "file_id": update.message.voice.file_id,
                    "file_size": update.message.voice.file_size,
                    "duration": update.message.voice.duration,
                },
            )
        except Exception:
            logger.debug("Failed to log voice message", exc_info=True)
        _schedule_background_task(
            context,
            process_audio_file(update, context, update.message.voice, beta_enabled=beta_enabled),
            description="voice_processing",
        )
        return

    # Обработка документов (видео/аудио файлы)
    if update.message.document:
        document = update.message.document
        filename = document.file_name.lower() if document.file_name else ""

        # Проверяем, является ли документ видео или аудио
        if any(ext in filename for ext in VIDEO_FORMATS):
            logger.info(f"Получен видео-документ от пользователя {user_id}: {filename}")
            try:
                log_telegram_event(
                    update.effective_user,
                    "message_document_video",
                    {
                        "chat_id": chat_id,
                        "file_id": document.file_id,
                        "file_size": document.file_size,
                        "file_name": document.file_name,
                    },
                )
            except Exception:
                logger.debug("Failed to log document video", exc_info=True)
            _schedule_background_task(
                context,
                process_video_file(update, context, document, beta_enabled=beta_enabled),
                description="message_document_video_processing",
            )
            return
        elif any(ext in filename for ext in AUDIO_FORMATS):
            logger.info(f"Получен аудио-документ от пользователя {user_id}: {filename}")
            try:
                log_telegram_event(
                    update.effective_user,
                    "message_document_audio",
                    {
                        "chat_id": chat_id,
                        "file_id": document.file_id,
                        "file_size": document.file_size,
                        "file_name": document.file_name,
                    },
                )
            except Exception:
                logger.debug("Failed to log document audio", exc_info=True)
            _schedule_background_task(
                context,
                process_audio_file(update, context, document, beta_enabled=beta_enabled),
                description="message_document_audio_processing",
            )
            return

    # Если это обычное текстовое сообщение
    if update.message.text:
        # Проверяем, что это не группа или что бот упомянут в сообщении
        is_group = update.effective_chat and update.effective_chat.type in ("group", "supergroup")
        bot_mentioned = False

        if is_group:
            # В группах проверяем, упомянут ли бот
            bot_username = context.bot.username
            if bot_username and f"@{bot_username}" in update.message.text:
                bot_mentioned = True

        # Отвечаем только в личных чатах или если бот упомянут в группе
        if not is_group or bot_mentioned:
            # Проверяем, ожидаем ли мы задачу для обработки транскрипции
            if context.user_data.get('waiting_for_task', False):
                await handle_transcript_processing_task(update, context)
            else:
                await update.message.reply_text(
                    "Привет! 🐱 Отправь мне видео или аудио файл, и я создам для тебя транскрипцию!\n\n"
                    "Поддерживаемые форматы:\n"
                    "📹 Видео: MP4, AVI, MOV, MKV и другие\n"
                    "🎵 Аудио: MP3, WAV, M4A, OGG и другие\n"
                    "🎤 Голосовые сообщения\n\n"
                    "Максимальный размер файла: 2 ГБ\n"
                    "Максимальная длительность: 4 часа\n\n"
                    "Используй /help для получения дополнительной информации!"
                )
            try:
                log_telegram_event(
                    update.effective_user,
                    "message_text",
                    {
                        "chat_id": chat_id,
                        "text": update.message.text,
                    },
                )
            except Exception:
                logger.debug("Failed to log text message", exc_info=True)

async def handle_transcript_processing_task(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обрабатывает текстовые сообщения с задачами для обработки транскрипции."""
    try:
        user_id = update.effective_user.id
        task_description = update.message.text

        # Сбрасываем флаг ожидания задачи
        context.user_data['waiting_for_task'] = False

        # Отправляем сообщение о начале обработки
        processing_msg = await update.message.reply_text(
            "🤖 Обрабатываю транскрипцию согласно твоей задаче...\n\n"
            "*сосредоточенно работает*\n"
            "Это может занять некоторое время...",
            parse_mode='Markdown'
        )

        # Получаем последнюю транскрипцию пользователя из базы данных
        from transkribator_modules.db.database import SessionLocal, UserService, TranscriptionService

        db = SessionLocal()
        try:
            user_service = UserService(db)
            transcription_service = TranscriptionService(db)

            # Получаем пользователя
            user = user_service.get_or_create_user(telegram_id=user_id)

            # Получаем последнюю транскрипцию пользователя
            transcriptions = transcription_service.get_user_transcriptions(user, limit=1)

            if not transcriptions:
                await processing_msg.edit_text(
                    "❌ Не найдено транскрипций для обработки.\n\n"
                    "Сначала отправьте файл для транскрипции!"
                )
                return

            latest_transcription = transcriptions[0]
            transcript_text = latest_transcription.formatted_transcript or latest_transcription.raw_transcript

            if not transcript_text:
                await processing_msg.edit_text("❌ Транскрипция пуста")
                return

            # Обрабатываем транскрипцию согласно задаче
            processed_text = await process_transcript_with_task(transcript_text, task_description)

            if not processed_text:
                await processing_msg.edit_text(
                    "❌ Не удалось обработать транскрипцию.\n\n"
                    "Возможно, сервис временно недоступен. Попробуйте позже."
                )
                return

            # Отправляем результат
            result_text = f"✅ **Результат обработки:**\n\n{processed_text}\n\n@CyberKitty19_bot"

            # Если результат длинный, отправляем файлом
            if len(result_text) > 4000:
                from docx import Document
                from pathlib import Path

                # Убеждаемся, что директория существует
                TRANSCRIPTIONS_DIR.mkdir(parents=True, exist_ok=True)
                docx_path = TRANSCRIPTIONS_DIR / f"processed_transcript_{user_id}.docx"

                document = Document()
                document.add_heading("Обработанная транскрипция", 0)
                document.add_paragraph(f"Задача: {task_description}")
                document.add_paragraph("Результат:")
                document.add_paragraph(processed_text)
                document.save(docx_path)

                with open(docx_path, 'rb') as f:
                    await update.message.reply_document(
                        document=f,
                        filename=f"processed_transcript.docx",
                        caption="✅ Результат обработки готов!\n\n@CyberKitty19_bot"
                    )

                # Удаляем временный файл
                docx_path.unlink(missing_ok=True)
            else:
                # Отправляем без parse_mode чтобы избежать ошибок с markdown entities
                await update.message.reply_text(result_text)

            # Обновляем сообщение о завершении
            await processing_msg.edit_text("✅ Обработка завершена!")

            # Показываем главное меню (личный кабинет)
            from transkribator_modules.bot.commands import personal_cabinet_command
            await personal_cabinet_command(update, context)

        finally:
            db.close()

    except Exception as e:
        logger.error(f"Ошибка при обработке задачи транскрипции: {e}")
        await update.message.reply_text(
            "❌ Произошла ошибка при обработке транскрипции.\n\n"
            "Попробуйте позже или обратитесь в поддержку."
        )

async def process_transcript_with_task(transcript_text: str, task_description: str) -> str:
    """Обрабатывает транскрипцию согласно задаче пользователя."""
    try:
        from transkribator_modules.transcribe.transcriber_v4 import request_llm_response

        system_prompt = (
            "Ты эксперт по обработке транскрипций. Твоя задача — читать запрос пользователя "
            "и выдавать готовый результат по предоставленной расшифровке.")
        user_prompt = (
            "ЗАДАЧА: {task}\n\n"
            "ТРАНСКРИПЦИЯ:\n{transcript}\n\n"
            "Обработай транскрипцию согласно задаче. Если указан конкретный формат, следуй ему точно."
        ).format(task=task_description, transcript=transcript_text)

        processed_text = None
        if request_llm_response:
            processed_text = await request_llm_response(system_prompt, user_prompt)

        if processed_text:
            cleaned_text = processed_text.strip().replace("*", "").replace("_", "").replace("`", "")
            if cleaned_text:
                return cleaned_text

        logger.warning("LLM не вернул результат для пользовательской задачи, отдаю исходную транскрипцию")
        return (
            "Не удалось обработать транскрипцию через ИИ. Вот исходная транскрипция:\n\n"
            f"{transcript_text}"
        )

    except Exception as e:
        logger.error(f"Ошибка при обработке транскрипции с задачей: {e}")
        return f"Произошла ошибка при обработке: {str(e)}"
